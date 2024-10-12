# * Imports
from flask import Blueprint, request, render_template, redirect, url_for, flash, session, jsonify,send_file
from sqlalchemy import inspect, Table,MetaData
from flask_login import login_required, current_user
from .models import User, Content
from . import db, Base
from .forms import AddSmartReportTemplateForm  # Import your form class
from io import BytesIO
from flask import send_file
import os
import tempfile
import shutil


# * Blueprint setup
app_admin_bp = Blueprint(
    'app_admin', __name__,
    static_folder='static',
    static_url_path='/static'
)

#todo: Global Error Handling
#@app_admin_bp.errorhandler(Exception)
#def handle_exception(e):
#    app_admin_bp.logger.error(f"Unhandled Exception in Admin Blueprint: {e}", exc_info=True)
#    return jsonify({'error': 'An internal error occurred'}), 500



# * Admin Dashboard Route
@app_admin_bp.route('/dashboard', methods=['GET', 'POST'])
@login_required
def admin_dashboard():
    if not current_user.is_admin:
        flash('Access restricted to administrators only.', 'warning')
        return redirect(url_for('user.login'))

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'add_contents':
            return redirect(url_for('app_admin.add_contents'))
        elif action == 'user_management':
            return redirect(url_for('app_admin.user_management'))
        elif action == 'reset_database':
            return redirect(url_for('app_admin.reset_database'))
        elif action == 'reset_users':
            return redirect(url_for('app_admin.reset_users'))

    users = db.session.query(User).all()
    contents = db.session.query(Content).all()
    
    return render_template('admin_dashboard.html', users=users, contents=contents)

#*----------------------------------------------------------------
# View all the models/tables in one place :
@app_admin_bp.route('/models')
@login_required
def view_models():
    if not current_user.is_admin:
        flash('Access restricted to administrators only.', 'warning')
        return redirect(url_for('main_routes.index'))
    
    # Use the Base metadata to get all model (table) names
    model_names = Base.metadata.tables.keys()
    tables_data = []
    for table_name in model_names:
        tables_data.append({
                'table_name': table_name,
                'endpoint': table_name
            })
    # Pass the list of tables data to the template
    return render_template('tables.html', tables_data=tables_data)
#*----------------------------------------------------------------
@app_admin_bp.route('/manage-model', methods=['GET', 'POST'])
@login_required
def manage_model():
    if not current_user.is_admin:
        flash('Access restricted to administrators only.', 'warning')
        return redirect(url_for('main_routes.index'))

    # Get the action and button ID from the form
    action = request.form.get('action')
    button_id = request.form.get('button_id')

    # Use the Base metadata to get all model (table) names
    model_names = Base.metadata.tables.keys()
    if action == "add_data" and button_id in model_names:
        # If action is 'add_data' and button ID matches a model name, redirect to the correct URL
        # Construct the URL to add new entries to the model via Flask-Admin
        target_url = f"/flask_admin/{button_id}/new/"
        return redirect(target_url)
    # ! Debugging 
    # If no specific action matches, render the test page with tables data and message in console
    return render_template(url_for('main_routes.debug'))
# *----------------------------------------------------------------
# create smart report templates:
# *----------------------------------------------------------------
@app_admin_bp.route('/create_template', methods=['GET', 'POST'])
@login_required
def create_report_template():
    form= AddSmartReportTemplateForm()
    return render_template('create_smart_report_template.html',form=form)

from io import BytesIO
import zipfile
from flask import send_file
import os

@app_admin_bp.route('/save_report_template', methods=['POST'])
@login_required
def save_report_template():
    form = AddSmartReportTemplateForm()
    
    if form.validate_on_submit():
        report_type = request.form.getlist('report_type')
        
        # Check if at least one report type is selected
        if not report_type:
            flash("Please select at least one report type (PDF or Word).", "error")
            return render_template('create_smart_report_template.html', form=form)
        
        # Prepare data dictionary from form fields
        data = {
            'template_name': form.template_name.data,
            'name': form.name.data,
            'gender': form.gender.data,
            'patient_id': form.patient_id.data,
            'age': form.age.data,
            'dob': form.dob.data,
            'location': form.location.data,
            'clinical_info': form.clinical_info.data,
            'technical_info': form.technical_info.data,
            'comparison': form.comparison.data,
            'observations': [{'section': obs.section.data, 'details': obs.details.data} for obs in form.observations],
            'conclusions': form.conclusions.data,
            'recommendations': form.recommendations.data
        }
        
        # Lists to collect file streams and filenames for zipping if needed
        files = []
        filenames = []
        
        # Generate the requested report types
        if 'word' in report_type:
            word_file = create_report_template_word(data, 'word', return_path_only=False)
            files.append(word_file)
            filenames.append('report_template.docx')
        
        if 'pdf' in report_type:
            pdf_file = create_report_template_pdf(data, return_path_only=False)
            files.append(pdf_file)
            filenames.append('report_template.pdf')
        
        # If both files were requested, zip them
        if len(files) == 2:
            zip_stream = BytesIO()
            zipf = zipfile.ZipFile(zip_stream, 'w', compression=zipfile.ZIP_DEFLATED)
            for file_stream, filename in zip(files, filenames):
                file_stream.seek(0)
                zipf.writestr(filename, file_stream.read())
            zipf.close()  # Explicitly close the ZipFile
            zip_stream.seek(0)
            return send_file(
            zip_stream,
            as_attachment=True,
            download_name='report_templates.zip',
            mimetype='application/zip')
        
        # If only one file is requested, send it directly
        elif len(files) == 1:
            file_stream = files[0]
            filename = filenames[0]
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if filename.endswith('.docx') else 'application/pdf'
            file_stream.seek(0)
            return send_file(file_stream, as_attachment=True, download_name=filename, mimetype=mime_type)
        
        else:
            flash("No report type was selected.", "error")
            return render_template('create_smart_report_template.html', form=form)
    
    # If form is not valid, re-render the form with errors
    return render_template('create_smart_report_template.html', form=form)
    
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from datetime import datetime


def create_report_template_word(data,report_type,return_path_only=False):
    # Initialize document
    doc = Document()
    
    # Hero container (unchanged)
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = title_paragraph.add_run()
    logo = logo_run.add_picture('app/static/assets/images/logo-white-bg.png', height=Inches(0.5))
    title_run = title_paragraph.add_run(" WSC - A Workstation Companion App")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(52, 58, 64)
    title_paragraph.space_after = Pt(6)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Simplify your radiology workflow with our intuitive and comprehensive tools.")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
    subtitle.space_after = Pt(2)
    
    italic_text = doc.add_paragraph()
    italic_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    italic_run = italic_text.add_run("Because every hero needs a trusty companion.")
    italic_run.italic = True
    italic_run.font.size = Pt(14)
    italic_run.font.color.rgb = RGBColor(40, 167, 69)
    italic_text.space_after = Pt(2)
    
    # Document Title (unchanged)
    title = doc.add_heading('Report Template', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 100, 0)
    template_name=data.get('template_name')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{template_name}")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
    subtitle.space_after = Pt(1)
    # Patient Information Table
    patient_info = {
        'Name': data.get('name', 'No name provided'),
        'Gender': data.get('gender', 'Not specified'),
        'ID': data.get('patient_id', 'No ID provided'),
        'Age': data.get('age', 'Not specified'),
        'DOB': data.get('dob', 'Not specified'),
        'Location': data.get('location', 'No location provided')
    }
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for idx, (field, value) in enumerate(patient_info.items()):
        row, col = divmod(idx, 2)
        cell = table.cell(row, col)
        cell.text = f"{field}: {value}"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.runs[0].font.size = Pt(12)
    
    # Clinical information section
    doc.add_heading('Clinical Context:', level=1)
    clinical_info_text = data.get('clinical_info', 'No clinical information provided')
    clinical_info_paragraph = doc.add_paragraph(clinical_info_text if clinical_info_text else 'No clinical information provided')
    clinical_info_paragraph.runs[0].font.size = Pt(12)
    clinical_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    clinical_info_paragraph.space_after = Pt(1)
    # technical information section
    doc.add_heading('Protocol/Technique:', level=1)
    technical_info_text = data.get('technical_info', 'No technique related infomration provided')
    technical_info_paragraph = doc.add_paragraph(technical_info_text if technical_info_text else 'No technique related infomration provided')
    technical_info_paragraph.runs[0].font.size = Pt(12)
    technical_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    technical_info_paragraph.space_after=Pt(1)
    #Comparison section
    doc.add_heading('Comparisons:', level=1)
    comparison_text = data.get('comparison', 'No relevant priors available for comparison')
    comparison_paragraph = doc.add_paragraph(comparison_text if comparison_text else 'No relevant priors available for comparison')
    comparison_paragraph.runs[0].font.size = Pt(12)
    comparison_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    comparison_paragraph
    # Observations Section
    doc.add_heading('Observations:', level=1)
    # Ensure observations is always a list, with a default entry if the list is empty or None
    observations = data.get('observations')
    section_name=""
    detail = ""
    for obs in observations:
        for key,value in obs.items():
            if key =='section':
                if value=="":
                    section_name='Section'
                else:
                    section_name = value
            
            if key=="details":
                if value=="":
                    detail= 'No details provided'
                else:
                    detail = value

        # Create a new paragraph for each observation
        obs_paragraph = doc.add_paragraph()
    
        # Section name in bold
        section_heading = obs_paragraph.add_run(f"{section_name}: ")
        section_heading.bold = True
        section_heading.font.size = Pt(13)
        section_heading.font.color.rgb = RGBColor(0, 102, 204)  # Light Blue
    
        # Details in regular font
        details_run = obs_paragraph.add_run(detail)
        details_run.font.size = Pt(12)
        obs_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Conclusions Section
    doc.add_heading('Conclusions:', level=1)
    conclusions_text = data.get('conclusions', 'No conclusions provided')
    conclusions_paragraph = doc.add_paragraph(conclusions_text if conclusions_text else 'No conclusions provided')
    conclusions_paragraph.runs[0].font.size = Pt(12)
    conclusions_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusions_paragraph
    # Recommendations Section
    doc.add_heading('Recommendations:', level=1)
    recommendations_text = data.get('recommendations', 'No recommendations provided')
    recommendations_paragraph = doc.add_paragraph(recommendations_text if recommendations_text else 'No recommendations provided')
    recommendations_paragraph.runs[0].font.size = Pt(12)
    recommendations_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Signature and Footer (unchanged)
    signature = doc.add_paragraph()
    signature_run = signature.add_run("\nReported and electronically signed by:\nRegistration Number:\n")
    signature_run.font.size = Pt(12)
    signature_run.bold = True
    date_run = signature.add_run(f"Dated: {datetime.now().strftime('%Y-%m-%d')}")
    date_run.font.size = Pt(12)
    date_run.italic = True

    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = 'SmartReportTemplates\nBrought to you by WSCompanion: because every Hero needs a companion'
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.runs[0].font.size = Pt(10)
    footer_paragraph.runs[0].font.color.rgb = RGBColor(107, 142, 35)
    
    # Use a temporary directory
    with tempfile.TemporaryDirectory() as temp_dir:
        word_file_path = os.path.join(temp_dir, "report_template.docx")
        doc.save(word_file_path)
        
        if return_path_only:
            # Since the temporary directory will be deleted after the 'with' block,
            # we need to copy the file if we need the path
            temp_word_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_word_file.close()
            shutil.copy(word_file_path, temp_word_file.name)
            return temp_word_file.name
        else:
            # Read the file into a BytesIO buffer
            with open(word_file_path, "rb") as word_file:
                file_stream = BytesIO(word_file.read())
            return file_stream


# Function to create pdf document form form data :from fpdf import FPDF
from reportlab.lib.pagesizes import LETTER, A4
from reportlab.lib.units import inch, mm
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def create_report_template_pdf(data, return_path_only=False):
    # Create a BytesIO buffer to hold the PDF data
    buffer = BytesIO()
    
    # Create a SimpleDocTemplate object
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Build the PDF elements
    elements = []
    
    # Add logo and title
    # Assuming the logo is in 'app/static/assets/images/logo-white-bg.png'
    logo_path = 'app/static/assets/images/logo-white-bg.png'
    try:
        logo = Image(logo_path, width=1*inch, height=1*inch)
        logo.hAlign = 'CENTER'
        elements.append(logo)
    except Exception as e:
        pass  # If logo not found, skip it
    
    title_style = ParagraphStyle(
        name='TitleStyle',
        fontSize=20,
        leading=24,
        alignment=1,  # Center alignment
        textColor=colors.HexColor('#343A40'),
        spaceAfter=6
    )
    title = Paragraph("WSC - A Workstation Companion App", title_style)
    elements.append(title)
    
    subtitle_style = ParagraphStyle(
        name='SubtitleStyle',
        fontSize=16,
        leading=18,
        alignment=1,
        textColor=colors.HexColor('#6C757D'),
        spaceAfter=2
    )
    subtitle = Paragraph(
        "Simplify your radiology workflow with our intuitive and comprehensive tools.",
        subtitle_style
    )
    elements.append(subtitle)
    
    italic_text_style = ParagraphStyle(
        name='ItalicTextStyle',
        fontSize=14,
        leading=12,
        alignment=1,
        textColor=colors.HexColor('#28A745'),
        spaceAfter=12,
        italic=True
    )
    italic_text = Paragraph("Because every hero needs a trusty companion.", italic_text_style)
    elements.append(italic_text)
    
    # Document Title
    report_title_style = ParagraphStyle(
        name='ReportTitleStyle',
        fontSize=20,
        leading=24,
        alignment=1,
        textColor=colors.HexColor('#006400'),
        spaceAfter=6
    )
    report_title = Paragraph("Report Template", report_title_style)
    elements.append(report_title)
    
    template_name = data.get('template_name', '')
    template_name_style = ParagraphStyle(
        name='TemplateNameStyle',
        fontSize=14,
        leading=18,
        alignment=1,
        textColor=colors.HexColor('#6C757D'),
        spaceAfter=12
    )
    template_name_paragraph = Paragraph(template_name, template_name_style)
    elements.append(template_name_paragraph)
    
    # Patient Information Table
    patient_info_data = [
        ['Name:', data.get('name', 'No name provided'), 'Gender:', data.get('gender', 'Not specified')],
        ['ID:', data.get('patient_id', 'No ID provided'), 'Age:', data.get('age', 'Not specified')],
        ['DOB:', data.get('dob', 'Not specified'), 'Location:', data.get('location', 'No location provided')]
    ]
    
    table_style = TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
    ])
    
    patient_table = Table(patient_info_data, colWidths=[1*inch, 2.5*inch, 1*inch, 2.5*inch])
    patient_table.setStyle(table_style)
    elements.append(patient_table)
    elements.append(Spacer(1, 12))
    
    # Sections: Clinical Context, Protocol/Technique, Comparisons, Observations, Conclusions, Recommendations
    def add_section(heading, content):
        combined_style = ParagraphStyle(
        name='CombinedStyle',
        fontSize=12,
        leading=15,
        textColor=colors.black,
        spaceAfter=12,
        spaceBefore=12
        )
        combined_text = f'<b>{heading}</b> {content if content else "No information provided"}'
        elements.append(Paragraph(combined_text, combined_style))
        
        
    # Clinical Context
    add_section('Clinical Context:', data.get('clinical_info', 'No clinical information provided'))
    
    # Protocol/Technique
    add_section('Protocol/Technique:', data.get('technical_info', 'No technique information provided'))
    
    # Comparisons
    add_section('Comparisons:', data.get('comparison', 'No relevant priors available for comparison'))
    
    # Observations
    # setting styles for observation fields :
    
    # Section name in bold and blue color
    section_style = ParagraphStyle(
        name='SectionStyle',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=4
        )
    detail_style = ParagraphStyle(
        name='DetailStyle',
        fontSize=12,
        leading=15,
        leftIndent=20,
        spaceAfter=6
        )
    heading_style = ParagraphStyle(
        name='HeadingStyle',
        fontSize=14,
        leading=20,
        textColor=colors.black,
        spaceAfter=6,
        spaceBefore=12
    )
    elements.append(Paragraph('Observations:', heading_style))
    
    # Ensure observations is always a list, with a default entry if the list is empty or None
    
    observations = data.get('observations')
    section_name=""
    detail = ""
    for obs in observations:
        for key,value in obs.items():
            if key =='section':
                if value=="":
                    section_name='Section'
                else:
                    section_name = value
            
            if key=="details":
                if value=="":
                    detail= 'No details provided'
                else:
                    detail = value
        section_paragraph = Paragraph(f"{section_name}:", section_style)
        elements.append(section_paragraph)
        # Details
        detail_paragraph = Paragraph(detail, detail_style)
        elements.append(detail_paragraph)
        
    # Conclusions
    add_section('Conclusions:', data.get('conclusions', 'No conclusions provided'))
    
    # Recommendations
    add_section('Recommendations:', data.get('recommendations', 'No recommendations provided'))
    
    # Signature and Date
    signature_style = ParagraphStyle(
        name='SignatureStyle',
        fontSize=12,
        leading=15,
        spaceAfter=12,
        spaceBefore=36
    )
    registration_style=ParagraphStyle(
        name='RegistrationStyle',
        fontSize=12,
        leading=15,
        spaceAfter=12,
        spaceBefore=8,
        italic=True
    )
    signature_text = f"Reported and electronically signed by:                  "
    registration_no="Registration Number:               "
    signature_date = f"Dated: {datetime.now().strftime('%Y-%m-%d')}"
    elements.append(Paragraph(signature_text, signature_style))
    elements.append(Paragraph(registration_no, registration_style))
    elements.append(Paragraph(signature_date, signature_style))
    
    # Footer
    def add_footer(canvas, doc):
        footer_text = 'SmartReportTemplates\nBrought to you by WSCompanion: because every Hero needs a companion'
        canvas.saveState()
        canvas.setFont('Helvetica', 10)
        canvas.setFillColor(colors.HexColor('#6B8E23'))
        width, height = doc.pagesize
        footer_lines = footer_text.split('\n')
        y = 15 * mm
        for line in footer_lines:
            canvas.drawCentredString(width / 2.0, y, line)
            y -= 12  # Move up for next line
        canvas.restoreState()
    
    # Build the PDF
    doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)
    
    buffer.seek(0)
    
    if return_path_only:
        # Save buffer to a temporary file
        temp_pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf_file.write(buffer.getvalue())
        temp_pdf_file.close()
        return temp_pdf_file.name
    else:
        # Return the buffer for direct download
        return buffer

#----------------------------------------------------------------

# todo: Placeholder route for adding contents
@app_admin_bp.route('/reset_database', methods=['GET', 'POST'])
@login_required
def reset_db():
    if not current_user.is_admin:
        flash('Access restricted to administrators only.', 'warning')
        return redirect(url_for('app_user.login'))
    
    action = request.form.get('action', '')
    button_id = request.form.get('button_id', '')

    if action == "reset_users" and button_id == "users":
        # Retrieve all users except admins
        users_to_delete = db.session.query(User).filter(User.is_admin == False).all()

        # Delete non-admin users
        for user in users_to_delete:
            db.session.delete(user)

        # Commit the session to trigger cascading
        try:
            db.session.commit()
            flash("All users except admins deleted.", 'warning')
        except Exception as e:
            db.session.rollback()
            flash(f"An error occurred while deleting users: {e}", 'danger')

        return redirect('/flask_admin/users')

    elif action == 'reset_db' and button_id == 'db':
        flash("The resetting of the database is not enabled. Please contact the owner of the App if you want to perform this action.", 'warning')
        return redirect(url_for('app_admin.view_models'))
#!----------------------------------------------------------------