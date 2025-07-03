#
from flask import Blueprint, request, render_template, redirect, url_for, flash, session, jsonify, send_file, current_app
from flask_login import login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_mail import Message
import os
import shutil
import json
from datetime import datetime, timedelta
from pytz import timezone, UTC
IST = timezone('Asia/Kolkata')




from app import db, mail
from config import Config, basedir, userdir
from app.models import (
    User, UserContentState, CategoryNames, ModuleNames,
    UserData, UserProfile, UserFeedback, UserReportTemplate,
    CPDLog, CPDActivityType, UserCPDState, InteractionTypeEnum
)
from app.forms import LoginForm, AddCPDLogForm
from app.util import (
    generate_password_reset_token, verify_password_reset_token,
    generate_otp_secret, generate_otp_token
)
# * Blueprint setup
app_user_bp = Blueprint(
    'app_user', __name__,
    static_folder='static',
    static_url_path='/static'
)

# ----------------------------------------------------------------
# * Route t degenerate qr code for 2factor authentication:
# ----------------------------------------------------------------
from .util import load_default_data,add_default_admin,add_default_contents,add_anonymous_user
import pyotp
import qrcode
import io
@app_user_bp.route('/qrcode')
@login_required
def qrcode_route():
    totp_secret = session.get('totp_secret')
    if not totp_secret:
        return 'TOTP secret not found. Please start the 2FA setup process.', 400

    # Use the user's email as the identifier
    user_email = current_user.email
    # Generate the provisioning URI
    # First create TOTP object:
    totp=pyotp.TOTP(totp_secret)
    # Generate the provision URI that will be used by authenticator app and hence provide all inormation to auth app:
    totp_uri = totp.provisioning_uri(name=user_email, issuer_name='WSCompanion')
    # Generate QR code image for this provisioning uri:
    img = qrcode.make(totp_uri)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png')


# ----------------------------------------------------------------
# * enale 2fa:
#---------------------------------------------------------------
@app_user_bp.route('/enable_2fa')
@login_required  # Ensure the user is logged in
def enable_2fa():
    # Generate and store the TOTP secret
    totp_secret = generate_otp_secret()
    session['totp_secret'] = totp_secret
    return render_template('enable_2fa.html', totp_secret=totp_secret)
# ----------------------------------------------------------------
# * enale 2fa:
#---------------------------------------------------------------
@app_user_bp.route('/disable_2fa', methods=['POST'])
@login_required
def disable_2fa():
    user = current_user  # Assuming you have `current_user` from Flask-Login
    user.totp_secret = None
    db.session.commit()
    flash('Two-Factor Authentication has been disabled.', 'success')
    return redirect(url_for('app_user.user_management'))


# ----------------------------------------------------------------
# * verify qr code route 
#----------------------------------------------------------------
@app_user_bp.route('/verify_2fa', methods=['POST'])
@login_required
def verify_2fa():
    token = request.form.get('token')
    totp_secret = session.get('totp_secret')

    if not totp_secret:
        flash('TOTP secret not found.', 'danger')
        return redirect(url_for('app_user.user_management'))

    totp = pyotp.TOTP(totp_secret)
    if totp.verify(token):
        # Save the totp_secret to the user's record
        user = current_user
        user.totp_secret = totp_secret
        db.session.commit()
        session.pop('totp_secret', None)
        flash('Two-Factor Authentication has been enabled.', 'success')
        return redirect(url_for('app_user.user_management'))
    else:
        flash('Invalid token. Please try again.', 'danger')
        return redirect(url_for('app_user.enable_2fa'))


#----------------------------------------------------------------
# * Login and Logout routes
#----------------------------------------------------------------

# User Login Route
@app_user_bp.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()  # Instantiate the form
    login_failed = False  # Default value for login_failed
    if form.validate_on_submit():  # Checks if the form is submitted and valid
        username = form.username.data.strip().lower()
        password = form.password.data.strip()
        remember = 'remember' in request.form  # Check if "Remember Me" was selected
        
        # Query the user using SQLAlchemy's session
        user = db.session.query(User).filter_by(username=username).first()
        if user and user.check_password(password):
            # Set the session to be permanent only if "Remember Me" is checked
            session.permanent = remember
            login_user(user, remember=remember)
            session['user_id'] = user.id  # Manually setting session data
            session.modified = True       # Ensure that session modifications are recognized
            user.status='active' 
            # Update UserData fields
            user_data = db.session.query(UserData).filter_by(user_id=user.id).first()
            if user_data:
                user_data.last_interaction =datetime.now(UTC)
                # Copy current login to last login
                if user_data.current_login:
                    user_data.last_login = user_data.current_login
                # Update current login to the new login time
                user_data.current_login = datetime.now(UTC)
                if user_data.login_count is None:
                    user_data.login_count = 0
                user_data.login_count += 1
                user_data.interaction_type = InteractionTypeEnum.LOGGED_IN  # Ensure this value aligns with your enum or model definition
                user_data.last_interaction=datetime.now(UTC)
                user.status = "active"
                db.session.commit()  # Save changes to the database
            
            print(f'Log in successful! Welcome back {user.username}!<hr style="color:yellow;">', 'success')
            return redirect(url_for('main_routes.index'))
        else:
            login_failed = True
            flash('Login failed.<br>Please check your username and password and try again.', 'danger')

    return render_template('login.html', form=form, login_failed=login_failed)
# Logout route
@app_user_bp.route('/logout')
@login_required
def logout():
    # Ensure the user is authenticated
    if not current_user.is_authenticated:
        flash("You must be logged in to log out!", "info")
        return redirect(url_for('app_user.login'))

    # Update user data before logging out
    user_id = current_user.id
    user = db.session.query(User).filter_by(id=user_id).first()
    user_data = db.session.query(UserData).filter_by(user_id=user_id).first()

    if user and user_data:
        # Ensure last_login is not None
        if user_data.last_login and user_data.session_start_time is not None:
            # Ensure last_login is timezone-aware
            if user_data.last_login and user_data.last_login.tzinfo is None:
                user_data.last_login = UTC.localize(user_data.last_login)
            if user_data.current_login and user_data.current_login.tzinfo is None:
                user_data.current_login = UTC.localize(user_data.current_login)
            # Calculate the time spent since last login
            time_spent = datetime.now(UTC) - user_data.current_login
            time_spent_in_minutes = time_spent.total_seconds()/60
            # Prevent negative time_spent
            if time_spent_in_minutes < 0:
                user_data.time_spent = 0
            else:
                user_data.time_spent += time_spent_in_minutes
        else:
            # If session_start_time is None, we can't calculate time_spent
            # Optionally, you can set time_spent to zero or handle it as needed
            user_data.time_spent += 0

        # Update user data
        user_data.last_interaction = datetime.now(UTC)
        user_data.interaction_type = InteractionTypeEnum.LOGGED_OUT  # Corrected typo from 'loged_out' to 'logged_out'
        user.status = "inactive"  # Set user status to inactive

        # Commit the changes to the database
        db.session.commit()

    # Log the user out
    logout_user()
    flash('You have been successfully logged out.', 'info')
    return redirect(url_for('app_user.login'))


# *----------------------------------------------------------------
# User Registration Route
@app_user_bp.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password'].strip()
        retyped_password = request.form['retyped_password'].strip()
        email = request.form['email'].strip()

        # Check if retyped and primary password match
        if password != retyped_password:
            flash('Passwords do not match. Please retype them correctly.', 'warning')
            return redirect(url_for('app_user.register'))

        # Check if the username or email already exists
        existing_user = db.session.query(User).filter_by(username=username).first()
        existing_email = db.session.query(User).filter_by(email=email).first()
        if existing_user:
            flash('Username already exists. Please choose a different one.', 'warning')
            return redirect(url_for('app_user.register'))
        elif existing_email:
            flash('Email already registered. Please choose a different one.', 'warning')
            return redirect(url_for('app_user.register'))

        if username and password and email:
            try:
                # Create and add the new user
                hashed_password = generate_password_hash(password, method='pbkdf2:sha256')
                new_user = User(username=username.lower(), password=hashed_password, email=email.lower())
                db.session.add(new_user)
                db.session.commit()

                # Define the paths for the dummy profile picture and the target location
                dummy_profile_pic_path = os.path.join(basedir,'app','static', 'assets','images','dummy_profile_pic.png')
                profile_pic_folder = os.path.join(userdir, f"{new_user.id}",'profile_pic')
                profile_pic_path = os.path.join(profile_pic_folder, 'dummy_profile_pic.png')

                # Create the target directory if it doesn't exist
                os.makedirs(profile_pic_folder, exist_ok=True)

                try:
                    # Copy the dummy profile picture to the target directory
                    shutil.copy(dummy_profile_pic_path, profile_pic_path)
                except Exception as e:
                    flash(f'Error copying the dummy profile picture: {str(e)}', 'danger')
                    print(f'Error copying the dummy profile picture: {str(e)}')

                # Initialize UserProfile with default values, using lists instead of JSON
                profile = UserProfile(
                    user_id=new_user.id,
                    profile_pic='dummy_profile_pic.png',  # This should just be the file name
                    profile_pic_path=profile_pic_path,  # This is the full path to where the file is stored
                    preferred_categories=','.join([category.value for category in CategoryNames]),  # Store as comma-separated string
                    preferred_modules=','.join([module.value for module in ModuleNames]),  # Store as comma-separated string
                    created_at=datetime.now(UTC),
                    updated_at=datetime.now(UTC)
                )
                db.session.add(profile)

                # Initialize UserData with baseline information
                user_data = UserData(
                    user_id=new_user.id,
                    interaction_type=InteractionTypeEnum.REGISTERED,
                    feedback=None,
                    content_rating=None,
                    time_spent=0,
                    last_interaction=datetime.now(UTC),
                    current_login=datetime.now(UTC),  # Assume the user has just logged in
                    last_login=None,
                    session_start_time=None,
                    login_count=0
                )
                db.session.add(user_data)

                # Initialize UserContentState with baseline information
                user_content_state = UserContentState(
                    user_id=new_user.id,
                    modified_filepath=None,
                    annotations=None,
                    created_at=datetime.now(UTC),
                    updated_at=datetime.now(UTC)
                )
                db.session.add(user_content_state)

                # Commit all changes to the database
                db.session.commit()

                # Clear the session and login the new user
                session.clear()  # Clears the existing session data to avoid conflicts
                login_user(new_user)  # Logs in the newly created user, ensuring the session reflects the correct user
                flash('Registration successful! Profile and related data initialized. Please log in.', 'success')
                return redirect(url_for('app_user.user_management'))

            except Exception as e:
                # Rollback any changes if an error occurs
                db.session.rollback()
                flash(f'Registration failed due to an error: {str(e)}', 'danger')
                print(f'Registration failed due to an error: {str(e)}')
                return redirect(url_for('app_user.register'))  # Corrected to redirect to the registration page

        else:
            flash('All fields are required.', 'danger')

    return render_template('register.html')
# *----------------------------------------------------------------
# Forget username / reset passwrod route:
@app_user_bp.route('/credential_management', methods=['GET', 'POST'])
def credential_manager():
    if request.method == 'GET':
        action=request.args.get('action')
        if action=="forgot_password":
            return render_template('credential_manager.html')
    else:
        action=request.form.get('action')
        input_email=request.form.get('email')
        input_email=input_email.lower()
        user=db.session.query(User).filter_by(email=input_email).first()
        if user:
            user_email=user.email
            if action=='reset_password':
                try:
                    # Generate the password reset token
                    token = generate_password_reset_token({'email': user_email})
                    
                    # Create the reset link
                    reset_link = url_for('app_user.change_password', token=token, _external=True, _scheme='https')
                    
                    # Prepare the email message for password reset
                    msg = Message('Account Recovery Email from WS Companion', recipients=[user_email])
                    msg.html = render_template('reset_password_email.html', reset_link=reset_link)
                    msg.body = render_template('reset_password_email.txt', reset_link=reset_link)
                    # Send the email
                    mail.send(msg)
                    
                    flash('A password reset link has been sent to your email successfully!<hr>This link is valid only for 10 minutes.', 'success')
                    return redirect(url_for('app_user.login'))
                except Exception as e:
                    print(f'Failed to send test email: {str(e)}')
                    flash(f'Failed to send test email: {str(e)}', 'danger')
                    return redirect(url_for('app_user.credential_manager'))
            else: # implies the action must be retrieve_username
                try:
                    # Prepare the email message for username retrieval
                    msg = Message('Account Recovery Email from WS Companion', recipients=[user_email])
                    msg.html = render_template('retrieve_username_email.html', user=user)
                    msg.body = render_template('retrieve_username_email.txt', user=user)
                    mail.send(msg)
                    flash('An account recovery email has been sent to your email address successfully!', 'success')
                    return redirect(url_for('app_user.login'))
                except Exception as e:
                    print(f'Failed to send test email: {str(e)}')
                    flash(f'Failed to send test email: {str(e)}', 'danger')
                    return redirect(url_for('app_user.credential_manager'))
        else:
            flash(f'No user found with this email address! <hr> Please enter the email assocaited with your account.', 'danger')
            return redirect(url_for('app_user.credential_manager'))
    return render_template('credential_manager.html')

# Change Password Route
@app_user_bp.route('/change_password/<token>', methods=['GET', 'POST'])
def change_password(token):
    # If POST request, process the form
    if request.method == 'POST':
        # Get form data
        password = request.form.get('password')
        retyped_password = request.form.get('retyped-password')
        
        # Validate passwords match
        if password != retyped_password:
            flash('Passwords do not match. Please try again.', 'danger')
            return render_template('change_password.html', token=token)
        
        # Verify token and update password
        user_data = verify_password_reset_token(token)
        if not user_data:
            flash('The reset link is invalid or expired. Please try again.', 'danger')
            return redirect(url_for('app_user.credential_manager'))
        
        # Assume user_data contains email
        user = db.session.query(User).filter_by(email=user_data.get('email')).first()
        if not user:
            flash('No account associated with this email.', 'danger')
            return redirect(url_for('app_user.credential_manager'))
        
        # Update the user's password securely
        user.set_password(password)  # This method should hash the password before saving
        db.session.commit()
        
        flash('Your password has been reset successfully! You can now log in with your new password.', 'success')
        return redirect(url_for('app_user.login'))
    
    # Render the change password form
    return render_template('change_password.html', token=token)
# ********************************
# Update password route:
@app_user_bp.route('/update_password', methods=['GET', 'POST'])
@login_required
def update_password():
    if request.method == 'POST':
        # Get form data
        new_password = request.form.get('new-password')
        retyped_password = request.form.get('retyped-password')
        
        if new_password:
            # Validate new passwords match
            if new_password != retyped_password:
                flash('Passwords do not match. Please try again.', 'danger')
                return render_template('user_management.html')
    
            # Check if new password is different from current
            if current_user.check_password(new_password):
                flash('You are attempting to reuse your old password. Please enter a new password.', 'danger')
                return render_template('user_management.html')
    
            # Update the password securely
            try:
                current_user.set_password(new_password)  # Hashing the password
                db.session.commit()
                flash('Your password has been updated successfully!', 'success')
                return redirect(url_for('app_user.user_management'))
            except Exception as e:
                db.session.rollback()
                flash(f'An error occurred while updating your password: {str(e)}', 'danger')
                return redirect(url_for('app_user.user_management'))
        else:
            flash('You did not enter any passwords.', 'danger')

    return render_template('user_management.html')

# Chanage email route:
@app_user_bp.route('/confirm_email')
@login_required
def confirm_email():
    token = request.args.get('token')
    if not token:
        flash('Invalid or missing token.', 'danger')
        return redirect(url_for('app_user.user_management'))
    token_data = verify_password_reset_token(token)
    if token_data and 'email' in token_data:
        new_email = token_data['email']
        # Update the user's email
        try:
            current_user.email = new_email
            db.session.commit()
            flash('Your email has been updated successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('An error occurred while updating your email. Please try again.', 'danger')
        return redirect(url_for('app_user.user_management'))
    else:
        flash('The confirmation link is invalid or has expired.', 'danger')
        return redirect(url_for('app_user.user_management'))

# ! Account recovery opetions:
@app_user_bp.route('/recover_account', methods=['GET', 'POST'])
def recover_account():
    if request.method == 'POST':
        recovery_option = request.form.get('recovery_option')
        if recovery_option == 'phone':
            phone = request.form.get('phone')
            user = db.session.query(User).filter_by(phone=phone).first()
            if not user:
                flash('No account associated with this phone number.', 'danger')
                return redirect(url_for('app_user.recover_account'))
            # Send a token via SMS using your SMS service (Twilio, etc.)
            token = generate_password_reset_token(user.email)
            send_sms(phone, f'Your account recovery token: {token}')
            flash('A recovery token has been sent to your phone.', 'success')
            return redirect(url_for('app_user.verify_recovery_token'))

        elif recovery_option == 'recovery_email':
            email = request.form.get('email')
            user = db.session.query(User).filter_by(email=email).first()
            if not user:
                flash('No account associated with this recovery email.', 'danger')
                return redirect(url_for('app_user.recover_account'))
            # Send a token via email using your mail extension
            token = generate_password_reset_token(user.email)
            send_email(user.email, 'Account Recovery', f'Your account recovery token: {token}')
            flash('A recovery token has been sent to your email.', 'success')
            return redirect(url_for('app_user.verify_recovery_token'))

        elif recovery_option == 'authenticator':
            # Assuming you have 2FA set up already
            if current_user.is_authenticated and current_user.two_factor_enabled:
                flash('You can use your authenticator app to verify and recover your account.', 'success')
                return redirect(url_for('app_user.verify_recovery_token'))
            else:
                flash('No authenticator app is set up for this account.', 'danger')
                return redirect(url_for('app_user.recover_account'))
    
    # Render the recovery form if GET request
    return render_template('recover_account.html')


# *----------------------------------------------------------------
# User Profile/Account Page and Related Routes
# *----------------------------------------------------------------
# Route to serve the user_management page
@app_user_bp.route('/account', methods=['GET'])
@login_required
def user_management(): 
    categories = CategoryNames  # Enum to populate the categories select box
    modules = ModuleNames  # Enum to populate the modules select box
    # Fetch the current user's profile
    user_profile = db.session.query(UserProfile).filter_by(user_id=current_user.id).first()
    try:
        if user_profile:
            if user_profile.profile_pic:
                # Construct the relative path for the profile picture
                profile_pic_rel_path = f"{current_user.id}/profile_pic/{user_profile.profile_pic}"

                # Generate URL to serve the profile picture using the `serve_user_data` route
                profile_pic_path = url_for('serve_user_data', filename=profile_pic_rel_path)
            else:
                # Use the default profile picture if `profile_pic` is missing
                profile_pic_path = url_for('static', filename='assets/images/logo-white-bg.pngg')
            
            print(f"Profile picture path: {profile_pic_path}")
            # Embed email verification status.
            return render_template('user_management.html', categories=categories, modules=modules, profile_pic_path=profile_pic_path)
    
        else:
            # Fall-back to the default profile picture if no user profile exists
            return render_template('user_management.html', categories=categories, modules=modules, profile_pic_path=url_for('static', filename='assets/images/logo-white-bg.png'))

    except Exception as e:
        # Handle any unexpected errors gracefully, log the error, and fall back to the default profile picture
        print(f"An error occurred: {e}")
        return render_template('user_management.html', categories=categories, modules=modules, profile_pic_path=url_for('static', filename='assets/images/logo-white-bg.png'))
# .###############################
# PROFILE MANAGEMENT ROUTES
# ===============================

@app_user_bp.route('/profile_manager', methods=['GET', 'POST'])
@login_required
def profile_manager():
    if request.method == 'POST':
        action = request.form.get('action')
        username = current_user.username
        user_id = current_user.id

        if action == 'update_profile_pic':
            
            # Handle profile picture upload logic here
            pic = request.files.get('profile_pic')  # Correctly get the file from the request
            if pic:
                filename = secure_filename(pic.filename)
                target_folder = os.path.join(basedir, userdir, f"{user_id}",'profile_pic')
                os.makedirs(target_folder, exist_ok=True)  # Create the directory if it doesn't exist
                profile_pic_path = os.path.join(target_folder, filename)

                try:
                    # Save the uploaded picture
                    pic.save(profile_pic_path)

                    # Update the user profile in the database
                    user = db.session.query(UserProfile).filter_by(user_id=current_user.id).first()
                    old_pic = user.profile_pic
                    old_pic_path = user.profile_pic_path
                    archived_folder = os.path.join(basedir, 'archived', f"{user_id}",'profile_pic')
                    os.makedirs(archived_folder, exist_ok=True)
                    if "dummy" in old_pic:
                        try:
                            os.remove(old_pic_path)
                        except FileNotFoundError:
                            pass  # Ignore the error if the file doesn't exist (e.g., it was already deleted)

                    elif old_pic and "dummy" not in old_pic:
                        shutil.move(old_pic_path, archived_folder)  # Archive the old profile picture

                    # Update the profile picture path in the user's profile
                    user.profile_pic = filename
                    user.profile_pic_path = profile_pic_path
                    user.updated_at = datetime.now(UTC)

                    # Update the user data
                    user_data = db.session.query(UserData).filter_by(user_id=current_user.id).first()
                    if user_data:
                        user_data.interaction_type = InteractionTypeEnum.UPDATED_PROFILE_PIC
                        user_data.last_interaction = datetime.now(UTC)

                    db.session.commit()  # Commit the changes to the database
                    flash('Profile picture uploaded successfully!', 'info')
                except Exception as e:
                    db.session.rollback()  # Rollback the session on error
                    flash(f'An error occurred while uploading the profile picture: {e}', 'danger')
                    return redirect(url_for('app_user.profile_manager'))

                # Pass target_folder and filename to the redirected route for rendering
                return redirect(url_for('app_user.user_management'))
            
            else:
                flash('No profile picture uploaded.', 'warning')
                return redirect(url_for('app_user.profile_manager'))

        elif action == 'update_username':
            requested_username = request.form.get('username')
            if requested_username:
                if requested_username == current_user.username:
                    flash('Did you enter your existing username by mistake? Please choose a different username.', 'warning')
                else:
                    # Check if the requested username is already taken
                    existing_user = db.session.query(User).filter_by(username=requested_username).first()
                    if existing_user:
                        flash('Username already taken. Please choose a different username.', 'warning')
                    else:
                        try:
                            current_user.username = requested_username.lower()
                            user_data = db.session.query(UserData).filter_by(user_id=current_user.id).first()
                            user_data.last_interaction = datetime.now(UTC)
                            user_data.interaction_type = InteractionTypeEnum.UPDATED_USERNAME
                            db.session.commit()  # Commit the changes to the database
                            flash('Username updated successfully!', 'info')
                            return redirect(url_for('app_user.user_management'))
                        except Exception as e:
                            db.session.rollback()  # Rollback the session on error
                            flash(f'An error occurred while updating the username: {e}', 'danger')
            else:
                flash('No username entered. Please enter a username.', 'warning')
            return redirect(url_for('app_user.profile_manager'))
    
        elif action == 'update_email':
            requested_email = request.form.get('email')
            # Check if the email is different
            if requested_email == current_user.email:
                flash('You have entered your current email. Please enter a new email.', 'warning')
            else:
                # Check if the email is already in use
                existing_user = db.session.query(User).filter_by(email=requested_email).first()
                if existing_user:
                    flash('This email is already in use. Please choose a different email.', 'warning')
                else:
                    # Generate token with the new email
                    token = generate_password_reset_token({'email': requested_email})
                    # Create the confirmation link
                    confirm_link = url_for('app_user.confirm_email', token=token, _external=True)
                    # Send confirmation email
                    print(f"message will be sent to {requested_email}")
                    msg = Message('Confirm Your Email Change', recipients=[requested_email])
                    msg.html = render_template('confirm_email.html', confirm_link=confirm_link)
                    msg.body = render_template('confirm_email.txt', confirm_link=confirm_link)
                    mail.send(msg)
                    flash('A confirmation email has been sent to your new email address. Please check your email. Please note that this link is valid only for 10 minutes.', 'info')
                    return redirect(url_for('app_user.user_management'))
                        
        elif action=="add_recovery_phone":
            recovery_phone = request.form.get('recovery_phone')
            if recovery_phone:
                try:
                    current_user.recovery_phone = recovery_phone
                    db.session.add(current_user)
                    db.session.commit()
                    db.session.commit()
                    user=db.session.query(User).filter_by(id=current_user.id).first()
                    user_email=user.email
                    user_name=user.username
                    account_update_type="Added Recovery Phone"
                    #send intimation to the user that a recovery phone has been added.
                    #create reset link to secure accunt :
                        # Generate the password reset token
                    token = generate_password_reset_token({'email': user_email})
                    
                    # Create the reset link
                    reset_link = url_for('app_user.change_password', token=token, _external=True, _scheme='https')
                    
                    msg=Message("Important: Suspicious activity on your ws-companion account", recipients=[user_email])
                    msg.html = render_template('account_updates_email.html',user_name=user_name,account_update_type=account_update_type,recovery_phone=recovery_phone,reset_link=reset_link)
                    msg.body = render_template('account_updates_email.txt', reset_link=reset_link)
                    mail.send(msg)
                    flash('Recovery phone number added successfully!', 'info')
                except Exception as e:
                    db.session.rollback()  # Rollback the session on error
                    flash(f'An error occurred while updating the recovery phone number: {e}', 'danger')
                    user=db.session.query(User).filter_by(id=current_user.id).first()
            else:
                flash('Please enter a valid recovery phone number')
            pass

    categories = CategoryNames
    modules = ModuleNames
    return render_template('user_management.html')

@app_user_bp.route('/delete_account')
def delete_account():
    flash('this will delete your account permanently')
    return ("This route of handle account deletion initiated by user")



# ===============================
# FINANCE MANAGEMENT ROUTES
# ===============================

@app_user_bp.route('/finance_manager', methods=['GET', 'POST'])
@login_required
def finance_manager():
    # Placeholder for finance-related actions
    if request.method == 'POST':
        # Handle finance management logic here
        flash('Finance management functionality will be implemented here.', 'info')
    return render_template('finance_manager.html')

# Removed individual finance routes like check_subscription as they are now covered by finance_manager.

# ===============================
# SECURITY MANAGEMENT ROUTES
# ===============================

@app_user_bp.route('/security_manager', methods=['GET', 'POST'])
@login_required
def security_manager():
    # Placeholder for security settings management (update questions, setup authenticator)
    if request.method == 'POST':
        # Handle security management logic here
        flash('Security management functionality will be implemented here.', 'info')
    return render_template('account_security.html')

# =======
# CPD dashboard mangment ROUTES
# =======
from app.models import UserCPDState, CPDLog
from dateutil.parser import parse as parse_date  # Add at the top

def safe_parse_date(text: str):
    try:
        return parse_date(text, fuzzy=True).date()
    except Exception:
        return None

import sqlalchemy as sa
@app_user_bp.route('/cpd/dashboard', methods=['GET', 'POST'])
@login_required
def cpd_dashboard():
    saved_cycles = (
        db.session.query(UserCPDState)
        .filter_by(user_id=current_user.id)
        .order_by(UserCPDState.created_at.desc())
        .all()
    )
    cpd_state = None
    appraisal_years = []
    active_year = None

    if request.method == 'GET' and request.args.get('new_cycle') == 'true':
        cpd_state = None

    elif request.method == 'GET':
        selected_cycle_id = request.args.get('cycle_id') or session.get('active_cpd_cycle_id')
        cpd_state = None
        if selected_cycle_id:
            cpd_state = db.session.get(UserCPDState, selected_cycle_id)
            if cpd_state:
                session['active_cpd_cycle_id'] = selected_cycle_id
            else:
                session.pop('active_cpd_cycle_id', None)
                flash("Previous appraisal cycle is no longer available. Please select a new one.", "warning")

    # Auto-select if only one saved cycle exists
    if not cpd_state and len(saved_cycles) == 1:
        cpd_state = saved_cycles[0]
        session['active_cpd_cycle_id'] = cpd_state.id
        if request.args.get('new_cycle') != 'true':  # Prevent redundant flash
            flash("✅ Only one appraisal cycle found. Activated automatically.", "info")

    if cpd_state and cpd_state.appraisal_cycle_start_date and cpd_state.appraisal_cycle_end_date:
        current = cpd_state.appraisal_cycle_start_date
        end = cpd_state.appraisal_cycle_end_date
        while current < end:
            next_year = current.replace(year=current.year + 1)
            if next_year > end:
                next_year = end
            appraisal_years.append({
                "start": current,
                "end": next_year,
                "label": f"{current.year}-{next_year.year}",
                "key": f"{current.year}-{next_year.year}"
            })
            current = next_year

    active_year_key = session.get('active_cpd_year_key')
    if active_year_key:
        for year in appraisal_years:
            if year['key'] == active_year_key:
                active_year = year
                break

    if request.method == 'POST' and request.args.get('new_cycle') == 'true':
        appraisal_start_input = request.form.get('appraisal_cycle_start')
        appraisal_end_input = request.form.get('appraisal_cycle_end')

        if appraisal_start_input and appraisal_end_input:
            try:
                appraisal_start_date = datetime.strptime(appraisal_start_input, '%Y-%m-%d').date()
                appraisal_end_date = datetime.strptime(appraisal_end_input, '%Y-%m-%d').date()

                appraisal_start_str = appraisal_start_date.strftime('%d/%m/%Y')
                appraisal_end_str = appraisal_end_date.strftime('%d/%m/%Y')

                start_year = appraisal_start_date.year
                end_year = appraisal_end_date.year

                existing_cycle = (
                    db.session.query(UserCPDState)
                    .filter_by(user_id=current_user.id)
                    .filter(
                        sa.extract('year', UserCPDState.appraisal_cycle_start_date) == start_year,
                        sa.extract('year', UserCPDState.appraisal_cycle_end_date) == end_year
                    )
                    .first()
                )

                if existing_cycle:
                    flash(f"⚠️ Appraisal cycle for {start_year} → {end_year} already exists.", "warning")
                    return redirect(url_for('app_user.cpd_dashboard', new_cycle='true'))

                appraisal_month = appraisal_start_date.strftime('%B')
                current_cpd_start = f"{appraisal_month} {datetime.now().year}"
                current_cpd_end = f"{appraisal_month} {datetime.now().year + 1}"

                new_state = UserCPDState(
                user_id=current_user.id,
                appraisal_cycle_start=appraisal_start_str,
                appraisal_cycle_end=appraisal_end_str,
                appraisal_cycle_start_date=appraisal_start_date,
                appraisal_cycle_end_date=appraisal_end_date,
                current_cpd_year_start=appraisal_start_date,
                current_cpd_year_end=appraisal_start_date.replace(year=appraisal_start_date.year + 1)
                )
                db.session.add(new_state)
                db.session.commit()
                session['active_cpd_cycle_id'] = new_state.id
                flash("✅ New appraisal cycle saved successfully.", "success")
                return redirect(url_for('app_user.cpd_dashboard'))

            except ValueError:
                flash("❌ Please enter valid dates in YYYY-MM-DD format.", "danger")

    # 🛠️ Core Fix: Fetch all logs for the user (not limited by cpd_state)
    cpd_logs = db.session.query(CPDLog).filter_by(user_id=current_user.id).all()

    current_year_logs, previous_year_logs = [], []
    total_current, total_previous = 0, 0

    if cpd_state and active_year:
        for log in cpd_logs:
            try:
                log_end = log.end_date.date()
                if active_year['start'] <= log_end <= active_year['end']:
                    total_current += (log.cpd_points_claimed or 0) + (1 if log.has_reflection else 0)
                    current_year_logs.append(log)
                else:
                    total_previous += (log.cpd_points_claimed or 0) + (1 if log.has_reflection else 0)
                    previous_year_logs.append(log)
            except Exception as e:
                print(f"⚠️ Error parsing log {log.id}: {e}")

    total_combined = total_current + total_previous
    remaining = max(0, 250 - total_combined)

    print("---- DEBUG INFO ----")
    print("Current user ID:", current_user.id)
    print("Selected CPD cycle:", cpd_state)
    print("Active appraisal year:", active_year)
    print("Current year logs count:", len(current_year_logs))
    print("Previous year logs count:", len(previous_year_logs))
    print("---------------------")

    return render_template(
        'cpd_dashboard.html',
        cpd_state=cpd_state,
        active_cycle=cpd_state,
        saved_cycles=saved_cycles,
        total_current=total_current,
        total_previous=total_previous,
        total_combined=total_combined,
        remaining=remaining,
        current_year_logs=current_year_logs,
        previous_year_logs=previous_year_logs,
        appraisal_years=appraisal_years,
        active_year=active_year ,
        new_cycle=request.args.get('new_cycle') == 'true'
    )

#====================
#This route allow user to add CPD activity #
#====================

import uuid
@app_user_bp.route("/app_user/cpd/add", methods=["GET", "POST"])
@login_required
def cpd_add():
    form = AddCPDLogForm()

    # ✅ Populate dropdown choices
    form.activity_type.choices = [
        (a.id, f"{a.name} ({a.default_credits})") 
        for a in db.session.query(CPDActivityType).all()
    ]

    if form.validate_on_submit():
        # === File upload handling ===
        filenames = []
        upload_folder = os.path.join(userdir, str(current_user.id), "cpd_certificates")
        os.makedirs(upload_folder, exist_ok=True)

        for file_field in form.certificate_files.entries:
            file = file_field.data
            if file and file.filename:
                ext = os.path.splitext(file.filename)[1].lower()
                unique_filename = f"{uuid.uuid4().hex}_{secure_filename(file.filename)}"
                path = os.path.join(upload_folder, unique_filename)
                file.save(path)
                filenames.append(unique_filename)

        # === Appraisal cycle state ===
        active_cycle_id = session.get("active_cpd_cycle_id")
        cpd_state = db.session.get(UserCPDState, active_cycle_id) if active_cycle_id else None

        if not cpd_state:
            flash("❌ You must select an active appraisal cycle before adding entries.", "danger")
            return redirect(url_for("app_user.cpd_dashboard"))

        start_date = form.start_date.data
        end_date = form.end_date.data

        if not form.title.data:
            flash("❌ Title is required.", "danger")
            return render_template("cpd_add.html", form=form)

        if not start_date:
            flash("❌ Start date is required.", "danger")
            return render_template("cpd_add.html", form=form)

        if not end_date:
            form.end_date.data = start_date
            form.end_date.raw_data = [start_date.strftime("%Y-%m-%d")]
            flash("ℹ️ End date was not entered. It has been set equal to the start date. Please confirm and resubmit.", "info")
            return render_template("cpd_add.html", form=form)

        active_year_key = session.get("active_cpd_year_key")
        active_year = None
        if active_year_key and cpd_state:
            start_yr, end_yr = map(int, active_year_key.split("-"))
            ay_start = cpd_state.appraisal_cycle_start_date.replace(year=start_yr)
            ay_end = ay_start.replace(year=start_yr + 1)
            active_year = {"start": ay_start, "end": ay_end}

        if not active_year:
            flash("❌ Could not determine active appraisal year. Please reselect the year.", "danger")
            return redirect(url_for("app_user.cpd_dashboard"))

        if start_date > active_year["end"]:
            flash("❌ Start date cannot be after Last date of this Appraisal year.", "danger")
            return render_template("cpd_add.html", form=form)

        if end_date < start_date:
            flash("❌ End date cannot be before start date.", "danger")
            return render_template("cpd_add.html", form=form)

        if end_date < active_year["start"] or end_date > active_year["end"]:
            flash(f"❌ End date must fall within the selected appraisal year: {active_year['start'].strftime('%d %b %Y')} to {active_year['end'].strftime('%d %b %Y')}", "danger")
            return render_template("cpd_add.html", form=form)
        # === Sanitize external link ===
        entered_external_links = form.external_links.data.strip()
        if entered_external_links and not entered_external_links.startswith(("http://", "https://")):
            external_links = "https://" + entered_external_links
        else:
            external_links = entered_external_links

        log = CPDLog(
            user_id=current_user.id,
            title=form.title.data,
            start_date=start_date,
            end_date=end_date,
            activity_type_id=form.activity_type.data,
            cpd_points_guideline=form.cpd_points_guideline.data,
            cpd_points_claimed=form.cpd_points_claimed.data,
            has_reflection=form.has_reflection.data,
            description=form.description.data,
            reflection=form.reflection.data,
            external_links=external_links,
            tags=form.tags.data,
            notes=form.notes.data,
            certificate_filenames=json.dumps(filenames),
            cpd_year_start=active_year["start"].strftime('%B %Y'),
            cpd_year_end=active_year["end"].strftime('%B %Y'),
            appraisal_cycle_start=cpd_state.appraisal_cycle_start,
            appraisal_cycle_end=cpd_state.appraisal_cycle_end,
            cpd_state_id=cpd_state.id
        )

        db.session.add(log)
        db.session.commit()
        flash("✅ CPD entry saved successfully", "success")
        return redirect(url_for("app_user.cpd_dashboard"))

    return render_template("cpd_add.html", form=form)

#====================
#This route allow user to edit CPD entry #

@app_user_bp.route('/cpd/edit/<uuid:log_id>', methods=['GET', 'POST'])
@login_required
def cpd_edit(log_id):
    log = db.session.get(CPDLog, log_id)
    if not log or log.user_id != current_user.id:
        flash("❌ CPD entry not found or access denied.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))

    form = AddCPDLogForm(obj=log)
    form.activity_type.choices = [
        (a.id, f"{a.name} ({a.default_credits})") 
        for a in db.session.query(CPDActivityType).all()
    ]
    form.activity_type.data = log.activity_type_id

    if request.method == "POST" and form.validate_on_submit():
        # Required fields check
        if not form.title.data:
            flash("❌ Title is required.", "danger")
            return render_template("cpd_add.html", form=form, edit_mode=True, log=log)

        if not form.start_date.data:
            flash("❌ Start date is required.", "danger")
            return render_template("cpd_add.html", form=form, edit_mode=True, log=log)

        if not form.end_date.data:
            flash("❌ End date is required.", "danger")
            return render_template("cpd_add.html", form=form, edit_mode=True, log=log)

        start_date = form.start_date.data
        end_date = form.end_date.data

        if end_date < start_date:
            flash("❌ End date cannot be before start date.", "danger")
            return render_template("cpd_add.html", form=form, edit_mode=True, log=log)

        # Appraisal year validation
        active_cycle_id = session.get("active_cpd_cycle_id")
        cpd_state = db.session.get(UserCPDState, active_cycle_id) if active_cycle_id else None

        active_year_key = session.get("active_cpd_year_key")
        active_year = None
        if active_year_key and cpd_state:
            start_yr, end_yr = map(int, active_year_key.split("-"))
            # Derive fixed active year boundaries
            ay_start = cpd_state.appraisal_cycle_start_date.replace(year=start_yr)
            ay_end = ay_start.replace(year=start_yr + 1)
            active_year = {"start": ay_start, "end": ay_end}

        if not active_year:
            flash("❌ Could not determine active appraisal year. Please reselect the year.", "danger")
            return redirect(url_for("app_user.cpd_dashboard"))

        if end_date < active_year["start"] or end_date > active_year["end"]:
            flash(f"❌ End date must fall within the selected appraisal year: {active_year['start'].strftime('%d %b %Y')} to {active_year['end'].strftime('%d %b %Y')}", "danger")
            return render_template("cpd_add.html", form=form, edit_mode=True, log=log)
        # === Sanitize external link ===
        entered_external_links = form.external_links.data.strip()
        if entered_external_links and not entered_external_links.startswith(("http://", "https://")):
            external_links = "https://" + entered_external_links
        else:
            external_links = entered_external_links

        # Update log entry
        log.title = form.title.data
        log.description = form.description.data
        log.reflection = form.reflection.data
        log.has_reflection = form.has_reflection.data
        log.cpd_points_guideline = form.cpd_points_guideline.data
        log.cpd_points_claimed = form.cpd_points_claimed.data
        log.start_date = start_date
        log.end_date = end_date
        log.tags = form.tags.data
        log.notes = form.notes.data
        log.external_links=external_links
        log.activity_type_id = int(form.activity_type.data)

        # Upload handling
        filenames = json.loads(log.certificate_filenames or "[]")
        upload_folder = os.path.join(userdir, str(current_user.id), "cpd_certificates")
        os.makedirs(upload_folder, exist_ok=True)

        for file_field in form.certificate_files.entries:
            file = file_field.data
            if file and file.filename:
                ext = os.path.splitext(file.filename)[1].lower()
                filename = secure_filename(file.filename)
                path = os.path.join(upload_folder, filename)
                file.save(path)
                filenames.append(filename)

        log.certificate_filenames = json.dumps(filenames)
        db.session.commit()
        flash("✅ CPD entry updated successfully!", "success")
        return redirect(url_for("app_user.cpd_dashboard"))

    return render_template("cpd_add.html", form=form, edit_mode=True, log=log)
#====================

#====================
#This route allow user to deleate Appraisal cylce #
#====================
from flask import request, flash, redirect, url_for
import uuid

@app_user_bp.route('/cpd/delete_cycle', methods=['POST'])
@login_required
def delete_cpd_cycle():
    cycle_id = request.form.get('cycle_id')
    if not cycle_id:
        flash("❌ No appraisal cycle selected. Please select a cycle to delete.", "warning")
        return redirect(url_for('app_user.cpd_dashboard'))

    try:
        cycle_uuid = uuid.UUID(cycle_id)
    except ValueError:
        flash("❌ Invalid cycle ID format.", "danger")
        return redirect(url_for('app_user.cpd_dashboard'))

    cycle_to_delete = db.session.get(UserCPDState, cycle_uuid)
    if not cycle_to_delete or cycle_to_delete.user_id != current_user.id:
        flash("❌ Invalid appraisal cycle or permission denied.", "danger")
        return redirect(url_for('app_user.cpd_dashboard'))

    try:
        logs_to_delete = db.session.query(CPDLog).filter_by(
            user_id=current_user.id,
            cpd_state_id=cycle_to_delete.id
        ).all()

        cert_dir = os.path.join(current_app.root_path, "user_data", str(current_user.id), "cpd_certificates")

        for log in logs_to_delete:
            cert_files = json.loads(log.certificate_filenames or "[]")
            for filename in cert_files:
                file_path = os.path.join(cert_dir, filename)
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        print(f"⚠️ Failed to delete certificate file {file_path}: {e}")

            db.session.delete(log)

        db.session.delete(cycle_to_delete)

        if session.get('active_cpd_cycle_id') == str(cycle_uuid):
            session.pop('active_cpd_cycle_id', None)

        db.session.commit()
        flash("✅ Appraisal cycle and all associated CPD data deleted successfully.", "success")

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error while deleting appraisal cycle: {e}")
        flash("❌ Failed to delete appraisal cycle. Please try again.", "danger")

    return redirect(url_for('app_user.cpd_dashboard'))
#===================
# route to set active year#
#================
@app_user_bp.route('/cpd/set_active_year/<year_key>', methods=['POST'])
@login_required
def set_active_cpd_year(year_key):
    session['active_cpd_year_key'] = year_key
    flash(f"✅ {year_key} was selected as currently active appraisal year.", "success")

    # Ensure the current cycle is passed back to rebuild the year cards
    active_cycle_id = session.get('active_cpd_cycle_id')
    return redirect(url_for('app_user.cpd_dashboard', cycle_id=str(active_cycle_id)))
#===================
# route to clear selected apprisal year#
#================
@app_user_bp.route('/clear_active_year', methods=['POST'])
@login_required
def clear_active_year():
    session.pop('active_cpd_year_key', None)
    flash("🔁 You can now select another appraisal year.", "info")
    
    active_cycle_id = session.get('active_cpd_cycle_id')
    if active_cycle_id:
        return redirect(url_for('app_user.cpd_dashboard', cycle_id=active_cycle_id))
    else:
        return redirect(url_for('app_user.cpd_dashboard'))
#===================
# route to set active appraisal cycle (for activation button)
#================
@app_user_bp.route('/cpd/set_active_cycle/<uuid:cycle_id>', methods=['POST'])
@login_required
def set_active_cpd_cycle(cycle_id):
    cpd_state = db.session.get(UserCPDState, cycle_id)
    if not cpd_state or cpd_state.user_id != current_user.id:
        flash("❌ Invalid or unauthorized appraisal cycle.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))

    session["active_cpd_cycle_id"] = str(cycle_id)
    session["active_cpd_year_key"] = None  # optional: reset year selection
    flash("✅ Appraisal cycle activated successfully.", "success")
    return redirect(url_for("app_user.cpd_dashboard", cycle_id=cycle_id))
#===============
#Route to serve CPD certificates.
#===========
import mimetypes
from flask import send_file, current_app, redirect, url_for, flash
from werkzeug.utils import safe_join
from config import userdir
@app_user_bp.route("/certificates/<filename>")
@login_required
def serve_certificate(filename):
    print(f"the filename is {filename}")
    cert_folder = os.path.join(userdir, str(current_user.id), "cpd_certificates")
    file_path = os.path.join(cert_folder, filename)
    print (f"{file_path} is filepath ")
    if not os.path.exists(file_path):
        flash("❌ Certificate not found.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))

    mimetype, _ = mimetypes.guess_type(file_path)
    if filename.lower().endswith('.pdf'):
        return send_file(file_path, mimetype=mimetype, as_attachment=False)
    else:
        return send_file(file_path, mimetype=mimetype or 'application/octet-stream', as_attachment=True)
#=====================
# This is  route for deleting cpd entry #
#====================
@app_user_bp.route('/cpd/delete/<uuid:log_id>', methods=['POST'])
@login_required
def delete_cpd_entry(log_id):
    log = db.session.get(CPDLog, log_id)
    
    if not log or log.user_id != current_user.id:
        flash("❌ Invalid CPD entry or permission denied.", "danger")
        return redirect(url_for('app_user.cpd_dashboard'))

    try:
        # Use correct absolute base directory
        from config import basedir
        cert_folder = os.path.join(basedir, "user_data", str(current_user.id), "cpd_certificates")

        # Delete uploaded certificate files
        if log.certificate_filenames:
            filenames = json.loads(log.certificate_filenames)
            for filename in filenames:
                path = os.path.join(cert_folder, filename)
                if os.path.exists(path):
                    os.remove(path)

        db.session.delete(log)
        db.session.commit()
        flash("✅ CPD entry and certificates deleted.", "success")

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error deleting CPD entry: {e}")
        flash("❌ Could not delete CPD entry.", "danger")

    return redirect(url_for('app_user.cpd_dashboard'))
#===================================================================#
# This is route for exproting the cpd log for single appraisal year #
#====================================================================#
@app_user_bp.route("/cpd/export_single_year", methods=["POST"])
@login_required
def export_single_year_log():
    from io import BytesIO
    from docx import Document
    from weasyprint import HTML

    export_format = request.form.get("export_format")
    full_name = request.form.get("full_name", current_user.username).strip()
    gmc_number = request.form.get("gmc_number", "").strip()
    active_cycle_id = session.get("active_cpd_cycle_id")
    active_year_key = session.get("active_cpd_year_key")

    if not active_cycle_id or not active_year_key:
        flash("❌ No active appraisal cycle or year selected.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))

    cpd_state = db.session.get(UserCPDState, active_cycle_id)
    if not cpd_state:
        flash("❌ Failed to load appraisal cycle.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))

    start_yr, _ = map(int, active_year_key.split("-"))
    ay_start = cpd_state.appraisal_cycle_start_date.replace(year=start_yr)
    ay_end = ay_start.replace(year=start_yr + 1)

    logs = (
        db.session.query(CPDLog)
        .filter(
            CPDLog.user_id == current_user.id,
            CPDLog.end_date.between(ay_start, ay_end)
        )
        .order_by(CPDLog.start_date.asc())
        .all()
    )

    if not logs:
        flash("No CPD logs found for this appraisal year.", "warning")
        return redirect(url_for("app_user.cpd_dashboard"))

    # Prepare data
    year_label = f"{ay_start.strftime('%B %Y')} → {ay_end.strftime('%B %Y')}"
    cpd_data = {year_label: []}
    yearly_totals = {year_label: 0}

    for log in logs:
        points = (log.cpd_points_claimed or 0) + (1 if log.has_reflection else 0)
        yearly_totals[year_label] += points
        cpd_data[year_label].append({
            "start_date": log.start_date.strftime('%d/%m/%Y'),
            "end_date": log.end_date.strftime('%d/%m/%Y'),
            "title": log.title,
            "claimed_points": points,
        })

    total_points = yearly_totals[year_label]
    points_deficit = max(0, 250 - total_points)

    context = {
        "cpd_state": cpd_state,
        "cpd_data": cpd_data,
        "yearly_totals": yearly_totals,
        "total_points": total_points,
        "points_deficit": points_deficit,
        "full_name": full_name,
        "gmc_number": gmc_number,
        "export_url": request.host_url.rstrip("/") + url_for("app_user.cpd_dashboard"),
        "export_scope": "single_year",
        "active_year": {"start": ay_start, "end": ay_end}
    }

    if export_format == "pdf":
        html = render_template("cpd_export_template.html", **context)
        pdf_io = BytesIO()
        HTML(string=html, base_url=request.host_url).write_pdf(pdf_io)
        pdf_io.seek(0)
        return send_file(pdf_io, mimetype="application/pdf", as_attachment=True,
                         download_name=f"CPD_Appraisal_Year_{start_yr}-{start_yr+1}.pdf")

    elif export_format == "word":
        doc = Document()
        doc.add_heading(f"CPD Appraisal Log: {year_label}", level=1)
        doc.add_paragraph(f"Name: {full_name}")
        if gmc_number:
            doc.add_paragraph(f"GMC Number: {gmc_number}")
        doc.add_paragraph("")

        doc.add_heading(year_label, level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Start Date"
        hdr_cells[1].text = "End Date"
        hdr_cells[2].text = "CPD Title"
        hdr_cells[3].text = "Claimed Points"

        for entry in cpd_data[year_label]:
            row_cells = table.add_row().cells
            row_cells[0].text = entry["start_date"]
            row_cells[1].text = entry["end_date"]
            row_cells[2].text = entry["title"]
            row_cells[3].text = str(entry["claimed_points"])

        doc.add_paragraph(f"Total Points This Year: {total_points}")
        doc.add_paragraph("")
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(f"Total Points: {total_points}")
        doc.add_paragraph(f"{'Deficit' if points_deficit > 0 else 'Excess'}: {abs(points_deficit)} CPD points from 250 requirement")

        word_io = BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        return send_file(
            word_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"CPD_Appraisal_Year_{start_yr}-{start_yr+1}.docx"
        )

    flash("❌ Invalid export format selected.", "danger")
    return redirect(url_for("app_user.cpd_dashboard"))

#===================
# route to export entire cpd log entire appraisal cycle #
#================
from flask import request, render_template, send_file, redirect, url_for, flash, session
from flask_login import login_required, current_user
from docx import Document
from weasyprint import HTML

@app_user_bp.route("/app_user/cpd/export_full_log", methods=["POST"])
@login_required
def export_full_appraisal_log():
    export_format = request.form.get("export_format")
    full_name = request.form.get("full_name", current_user.username).strip()
    if not full_name:
        flash("❌ Please enter your full name to proceed with the export.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))
    gmc_number = request.form.get("gmc_number", "").strip()
    active_cycle_id = session.get('active_cpd_cycle_id')

    if not active_cycle_id:
        flash("⚠️ No active appraisal cycle selected. Please choose a cycle before exporting.", "warning")
        return redirect(url_for("app_user.cpd_dashboard"))

    active_cycle = db.session.get(UserCPDState, active_cycle_id)
    if not active_cycle:
        flash("❌ Could not load the selected appraisal cycle.", "danger")
        return redirect(url_for("app_user.cpd_dashboard"))

    logs = (
        db.session.query(CPDLog)
        .filter_by(user_id=current_user.id,
                   appraisal_cycle_start=active_cycle.appraisal_cycle_start,
                   appraisal_cycle_end=active_cycle.appraisal_cycle_end)
        .order_by(CPDLog.start_date.asc())
        .all()
    )

    # Collect all unique CPD year keys across the entire cycle
    all_year_keys = set()
    for log in logs:
        year_key = f"{log.cpd_year_start} → {log.cpd_year_end}"
        all_year_keys.add(year_key)

    all_year_keys = sorted(
        all_year_keys,
        key=lambda k: datetime.strptime(k.split("→")[0].strip(), "%B %Y")
    )

    # Organize logs into year groups and compute totals
    cpd_data = {}
    yearly_totals = {}
    for year_key in all_year_keys:
        cpd_data[year_key] = []
        yearly_totals[year_key] = 0

    for log in logs:
        year_key = f"{log.cpd_year_start} → {log.cpd_year_end}"
        points = (log.cpd_points_claimed or 0) + (1 if log.has_reflection else 0)
        yearly_totals[year_key] += points

        cpd_data[year_key].append({
            "start_date": log.start_date.strftime('%d/%m/%Y') if log.start_date else "",
            "end_date": log.end_date.strftime('%d/%m/%Y') if log.end_date else "",
            "title": log.title,
            "claimed_points": points,
        })

    total_points = sum(yearly_totals.values())
    points_deficit = max(0, 250 - total_points)

    context = {
        "cpd_state": active_cycle,
        "cpd_data": cpd_data,
        "yearly_totals": yearly_totals,
        "total_points": total_points,
        "points_deficit": points_deficit,
        "full_name": full_name,
        "gmc_number": gmc_number,
        "export_url": request.host_url.rstrip("/") + url_for("app_user.cpd_dashboard"),
        "export_scope": "full_cycle"
    }

    if export_format == "pdf":
        html = render_template("cpd_export_template.html", **context)
        pdf_io = BytesIO()
        HTML(string=html, base_url=request.host_url).write_pdf(pdf_io)
        pdf_io.seek(0)
        return send_file(pdf_io, mimetype="application/pdf", as_attachment=True,
                        download_name="CPD_Appraisal_Log.pdf")

    elif export_format == "word":
        doc = Document()
        doc.add_heading(
            f"CPD Appraisal Log: {active_cycle.appraisal_cycle_start_date.strftime('%d/%m/%Y')} → {active_cycle.appraisal_cycle_end_date.strftime('%d/%m/%Y')}",
            level=1,
        )

        doc.add_paragraph(f"Name: {full_name}")
        if gmc_number:
            doc.add_paragraph(f"GMC Number: {gmc_number}")
        doc.add_paragraph("")

        for year in all_year_keys:
            doc.add_heading(year, level=2)
            entries = cpd_data.get(year, [])

            if not entries:
                doc.add_paragraph("No CPD entries recorded for this year.")
            else:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Start Date"
                hdr_cells[1].text = "End Date"
                hdr_cells[2].text = "CPD Title"
                hdr_cells[3].text = "Claimed Points"

                for entry in entries:
                    row_cells = table.add_row().cells
                    row_cells[0].text = entry["start_date"]
                    row_cells[1].text = entry["end_date"]
                    row_cells[2].text = entry["title"]
                    row_cells[3].text = str(entry["claimed_points"])

                doc.add_paragraph(f"Total Points This Year: {yearly_totals[year]}")

            doc.add_paragraph("")

        doc.add_paragraph("Summary", style="Heading 2")
        doc.add_paragraph(f"Total Points Across Cycle: {total_points}")
        doc.add_paragraph(f"{'Deficit' if points_deficit > 0 else 'Excess'}: {abs(points_deficit)} CPD points from 250 requirement")

        word_io = BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        return send_file(
            word_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="CPD_Appraisal_Log.docx"
        )

    flash("❌ Invalid export format selected.", "danger")
    return redirect(url_for("app_user.cpd_dashboard"))



# -------------------------
# Productivity Dashboard
# -------------------------
from datetime import datetime, timedelta
from pytz import UTC
from flask import render_template
from flask_login import login_required, current_user

@app_user_bp.route("/productivity/dashboard", methods=["GET"])
@login_required
def productivity_dashboard():
    user_profile = db.session.query(UserProfile).filter_by(user_id=current_user.id).first()

    def format_minutes(minutes):
        hrs = int(minutes // 60)
        mins = minutes % 60
        return f"{hrs}h {mins:.1f}m" if hrs else f"{mins:.1f}m"

    now_utc = datetime.now(UTC)
    today = now_utc.date()
    today_start = datetime.combine(today, datetime.min.time()).replace(tzinfo=UTC)
    today_end = datetime.combine(today, datetime.max.time()).replace(tzinfo=UTC)
    
    yesterday_start = today_start - timedelta(days=1)
    yesterday_end = today_end - timedelta(days=1)

    # Yesterday's logs
    yesterday_logs = (
        db.session.query(UserData)
        .filter(
            UserData.user_id == current_user.id,
            UserData.is_productivity_log == True,
            UserData.session_start_time >= yesterday_start,
            UserData.session_start_time <= yesterday_end,
        ).all()
    )

    yesterday_total_cases = sum(log.num_cases_reported or 0 for log in yesterday_logs)
    yesterday_total_minutes = sum(log.time_spent or 0 for log in yesterday_logs)
    yesterday_avg_minutes = (yesterday_total_minutes / yesterday_total_cases if yesterday_total_cases else 0)
    yesterday_total_time = format_minutes(yesterday_total_minutes)
    yesterday_avg_time_per_case = format_minutes(yesterday_avg_minutes) if yesterday_total_cases else "—"

    # Today's logs
    today_logs = (
        db.session.query(UserData)
        .filter(
            UserData.user_id == current_user.id,
            UserData.is_productivity_log == True,
            UserData.session_start_time >= today_start,
            UserData.session_start_time <= today_end,
        ).all()
    )

    total_cases = sum(log.num_cases_reported or 0 for log in today_logs)
    total_minutes = sum(log.time_spent or 0 for log in today_logs)
    avg_minutes = total_minutes / total_cases if total_cases else 0
    total_time = format_minutes(total_minutes)
    avg_time_per_case = format_minutes(avg_minutes) if total_cases else "—"
    
    return render_template(
        "productivity_dashboard.html",
        modules=list(ModuleNames),
        today_logs=today_logs,
        total_cases=total_cases,
        total_time=total_time,
        avg_time_per_case=avg_time_per_case,
        user_profile=user_profile,
        yesterday_logs=yesterday_logs,
        yesterday_total_cases=yesterday_total_cases,
        yesterday_total_time=yesterday_total_time,
        yesterday_avg_time_per_case=yesterday_avg_time_per_case,
        today_start=today_start,
        today_end=today_end
    )
#Save user productivty dashboard preferences :
@app_user_bp.route("/save_productivity_preferences", methods=["POST"])
@login_required
def save_productivity_preferences():
    try:
        # Fetch or create user profile
        profile = db.session.query(UserProfile).filter_by(user_id=current_user.id).first()
        if not profile:
            profile = UserProfile(user_id=current_user.id)

        # Get selected subspecialties (as strings) from form
        raw_subspecialties = request.form.getlist("preferred_subspecialties")
        enum_subspecialties = [ModuleNames(sub) for sub in raw_subspecialties]

        # Get workplaces string and convert to list
        raw_workplaces = request.form.get("preferred_workplaces", "")
        workplace_list = [w.strip() for w in raw_workplaces.split(",") if w.strip()]

        # Update profile fields
        profile.preferred_subspecialties = enum_subspecialties
        profile.preferred_workplaces = workplace_list
        profile.updated_at = datetime.now(UTC)

        # Add and commit
        db.session.add(profile)
        db.session.commit()
        flash("Preferences updated successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error saving preferences: {str(e)}", "danger")

    return redirect(url_for("app_user.productivity_dashboard"))

#captures submitted batch info and stores them in UserData.
from pytz import UTC
from datetime import datetime
from flask import request, redirect, url_for, flash
from flask_login import current_user
from dateutil import parser
@app_user_bp.route('/save_session_log', methods=['POST'])
def save_session_log():
    print("save session route reached")
    try:
        session_start_time_str = request.form.get("session_start_time")
        session_end_time_str = request.form.get("session_end_time")
        time_spent_str = request.form.get("time_spent")

        if not (session_start_time_str and session_end_time_str and time_spent_str):
            flash("Missing session data. Please try again.", "danger")
            return redirect(url_for('app_user.productivity_dashboard'))

        # Parse ISO strings (may or may not have tzinfo)
        
        start_dt = parser.isoparse(session_start_time_str)
        end_dt = parser.isoparse(session_end_time_str)

        # Force conversion to UTC
        start_dt = start_dt.astimezone(UTC)
        end_dt = end_dt.astimezone(UTC)

        session_start_time = start_dt
        session_end_time = end_dt
        time_spent = int(time_spent_str)

        cases = request.form.getlist("cases[]")
        modalities = request.form.getlist("modalities[]")
        workplaces = request.form.getlist("workplaces[]")
        notes_list = request.form.getlist("notes[]")

        user_id = current_user.id if current_user.is_authenticated else add_anonymous_user()

        for i in range(len(cases)):
            log = UserData(
                session_start_time=session_start_time,
                session_end_time=session_end_time,
                time_spent=time_spent if time_spent else None,
                num_cases_reported=int(cases[i]) if cases[i] else 0,
                modalities_handled=[modalities[i]] if modalities[i] else [],
                session_type=workplaces[i],
                notes=notes_list[i],
                user_id=user_id,
                is_productivity_log=True
            )
            db.session.add(log)

        db.session.commit()
        flash("✅ Session log saved successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error saving session log: {str(e)}", "danger")

    return redirect(url_for('app_user.productivity_dashboard'))