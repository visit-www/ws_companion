#Imports
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from flask import current_app
import json
import re
from datetime import datetime, timezone
from .models import (
    db,
    Content,
    ImagingProtocol,
    ClassificationSystem,
    User,
    UserData,
    UserContentState,
    AdminReportTemplate,
    UserReportTemplate,
    ClassificationCategoryEnum,
    ModuleNames,
    BodyPartEnum,
    ModalityEnum,
    CategoryNames,
    NormalMeasurement,
)
from sqlalchemy import or_
import os
from config import basedir,ADMIN_EMAIL, ADMIN_PASSWORD,ANONYMOUS_EMAIL,ANONYMOUS_PASSWORD,ANONYMOUS_USER_ID
# *-------------------------------------------------------------------------

def generate_password_reset_token(data, expiration=600):
    """Generate a secure token for given data."""
    serializer = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    return serializer.dumps(data, salt=current_app.config['SECURITY_PASSWORD_SALT'])

def verify_password_reset_token(token, expiration=600):
    """Verify a password reset token and return the associated data if valid."""
    serializer = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    try:
        data = serializer.loads(token, salt=current_app.config['SECURITY_PASSWORD_SALT'], max_age=expiration)
        return data
    except (BadSignature, SignatureExpired):
        return None

#Functions to generate otps using pyotp for mfa:

import pyotp
def generate_otp_secret():
    #Generates a new TOTP secret key
    return pyotp.random_base32()

def generate_otp_token(secret, interval=200):
    #Generates a TOTP token based on the given secret and interval
    totp = pyotp.TOTP(secret, interval=interval)
    return totp.now()

# *-----------------------------------------------------------
# Default app initialization
from sqlalchemy import or_
from config import ADMIN_EMAIL, ADMIN_PASSWORD

def add_default_admin(admin_data: dict) -> None:
    """
    Ensure there is at least one admin user.
    If an admin already exists (same email, same username, or is_admin=True),
    update its email/password to match current config.
    Otherwise, create a new admin user.
    """
    try:
        admin_user = (
            db.session.query(User)
            .filter(
                or_(
                    User.email == ADMIN_EMAIL,
                    User.username == admin_data.get("username", "admin"),
                    User.is_admin.is_(True),
                )
            )
            .first()
        )

        if admin_user:
            changed = False

            desired_username = admin_data.get("username", "admin")
            if admin_user.username != desired_username:
                admin_user.username = desired_username
                changed = True

            if admin_user.email != ADMIN_EMAIL:
                admin_user.email = ADMIN_EMAIL
                changed = True

            if not admin_user.is_admin:
                admin_user.is_admin = True
                changed = True

            if ADMIN_PASSWORD:
                admin_user.set_password(ADMIN_PASSWORD)
                changed = True

            if changed:
                db.session.commit()
                print(
                    f"Updated existing admin user: username={admin_user.username}, email={admin_user.email}"
                )
            else:
                print(
                    f"Admin user already up to date: username={admin_user.username}, email={admin_user.email}"
                )
        else:
            new_admin = User(
                username=admin_data.get("username", "admin"),
                email=ADMIN_EMAIL,
                is_paid=admin_data.get("is_paid", True),
                is_admin=admin_data.get("is_admin", True),
            )
            new_admin.set_password(ADMIN_PASSWORD)
            db.session.add(new_admin)
            db.session.commit()
            print(
                f"Admin user created: username={new_admin.username}, email={new_admin.email}"
            )

    except Exception as e:
        db.session.rollback()
        print(f"Error while adding/updating admin user: {e}")
        
# Crreate Anonymous user to relate to orphaned data after users or content is delated (referecnes, userfeedback)
from config import ANONYMOUS_USER_ID
def add_anonymous_user():
    anonymous_user = db.session.query(User).filter_by(username='anonymous').first()
    try:
        if not anonymous_user:
            # Create anonymous user
            anonymous_user = User(
                id=ANONYMOUS_USER_ID,
                username='anonymous',
                email=ANONYMOUS_EMAIL,
                is_paid=False,
                is_admin=False,
                status='active',
            )
            anonymous_user.set_password(ANONYMOUS_PASSWORD)
            db.session.add(anonymous_user)
            db.session.commit()
            print(f"Anonymous user created: {anonymous_user}, {anonymous_user.email}")
        else:
            print(f"Anonymous user already exists: {anonymous_user.username}")
    except Exception as e:
        print(f"Error adding anonymous user: {e}")
        db.session.rollback()
        pass

def add_default_contents(contents_data):
    """Add default contents if not already present."""
    for content_data in contents_data:
        existing_content = db.session.query(Content).filter_by(title=content_data['title']).first()
        if not existing_content:
            new_content = Content(
                title=content_data['title'],
                category=content_data['category'],
                module=content_data['module'],
                status=content_data['status'],
                external_url=content_data.get('external_url'),
                embed_code=content_data['embed_code'],
                description=content_data['description'],
                created_by=content_data['created_by'],
                language=content_data['language'],
                created_at=datetime.now(timezone.utc),
                updated_at=datetime.now(timezone.utc)
            )
            db.session.add(new_content)
            print(f"Content added: {new_content.title}")
        else:
            print(f"Content already exists: {existing_content.title}")
    

    db.session.commit()
    print("Default contents loaded successfully.")

# Utility function to add default admin templates
def add_default_admin_templates():
    """
    Seed a small, curated set of reporting templates into AdminReportTemplate.
    Idempotent: if a template with the same name already exists, it is skipped.
    """
    default_templates = [
        {
            "template_name": "MRI Brain – General Reporting Template",
            "modality": ModalityEnum.MRI,
            "body_part": BodyPartEnum.NEURO,
            "module": ModuleNames.NEURORADIOLOGY,
            "category": CategoryNames.REPORT_TEMPLATE,
            "tags": "mri, brain, neuro, routine",
            "file": None,
            "filepath": None,
        },
        {
            "template_name": "CT Chest – Pulmonary Embolism (CTPA)",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.LUNG,
            "module": ModuleNames.CHEST,
            "category": CategoryNames.REPORT_TEMPLATE,
            "tags": "ct, chest, ctpa, pulmonary embolism",
            "file": None,
            "filepath": None,
        },
        {
            "template_name": "CT Abdomen/Pelvis – Acute Abdomen",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.ABDOMEN,
            "module": ModuleNames.GASTROINTESTINAL,
            "category": CategoryNames.REPORT_TEMPLATE,
            "tags": "ct, abdomen, pelvis, acute abdomen, emergency",
            "file": None,
            "filepath": None,
        },
        {
            "template_name": "X-ray Chest – General / Preoperative",
            "modality": ModalityEnum.X_RAY,
            "body_part": BodyPartEnum.LUNG,
            "module": ModuleNames.CHEST,
            "category": CategoryNames.REPORT_TEMPLATE,
            "tags": "xray, chest, preop, general",
            "file": None,
            "filepath": None,
        },
    ]

    created = 0

    for tpl in default_templates:
        existing = db.session.query(AdminReportTemplate).filter_by(
            template_name=tpl["template_name"]
        ).first()
        if existing:
            continue

        new_tpl = AdminReportTemplate(
            template_name=tpl["template_name"],
            modality=tpl["modality"],
            body_part=tpl["body_part"],
            module=tpl["module"],
            category=tpl["category"],
            tags=tpl["tags"],
            file=tpl["file"],
            filepath=tpl["filepath"],
        )
        db.session.add(new_tpl)
        created += 1

    if created > 0:
        db.session.commit()
        print(f"Admin templates seeded: {created}")
    else:
        print("Admin templates already present, no new templates added.")

# Internal helper to clean free-text search queries for case-aware search
def _clean_search_text(raw_text: str | None):
    """
    Normalise a free-text search string by:
      - tokenising with _tokenize
      - removing very common stopwords (e.g. 'is', 'there', 'the', etc.)
    Returns a space-joined string of remaining tokens, or an empty string if none.
    """
    if not raw_text:
        return ""

    tokens = _tokenize(raw_text)
    # Very small, conservative stopword list aimed at question phrases
    stopwords = {
        "is",
        "are",
        "am",
        "was",
        "were",
        "be",
        "being",
        "been",
        "the",
        "a",
        "an",
        "of",
        "in",
        "on",
        "at",
        "to",
        "for",
        "with",
        "and",
        "or",
        "if",
        "any",
        "what",
        "why",
        "how",
        "there",
        "this",
        "that",
        "these",
        "those",
        "does",
        "do",
        "did",
        "has",
        "have",
        "had",
        "cause",
        "causes",
    }
    filtered = [t for t in tokens if t not in stopwords]
    if not filtered:
        return ""
    return " ".join(filtered)

# Shared intelligent body-part matcher for Case Workspace (protocols, templates, measurements, classifications)
def _body_part_matches(proto_bp, selected_bp):
    """Intelligent body-part matching for Case Workspace.

    Rules:
    - If no selected body part, everything passes.
    - Exact enum equality always passes.
    - OTHERS acts as a broad wildcard:
      * An entry tagged with OTHERS can appear for any selected body part.
      * If the selected body part is OTHERS, any body part is allowed.
    - ABDOMEN behaves as a superset of several abdominal regions:
      * ABDOMEN, UPPER_GI, LOWER_GI, HEPATOBILIARY, UROLOGY, GYNAECOLOGY.
    - LUNG / CARDIAC / VASCULAR form a thoracic family.
    - NEURO / HEAD_AND_NECK / VASCULAR form a neuro–head–neck family.
    """
    if selected_bp is None:
        return True
    if proto_bp is None:
        return False
    if proto_bp == selected_bp:
        return True

    try:
        proto_name = proto_bp.name.upper()
        sel_name = selected_bp.name.upper()
    except AttributeError:
        return False

    # Broad wildcard
    wildcard_set = {"OTHERS"}
    if proto_name in wildcard_set or sel_name in wildcard_set:
        return True

    # Abdomen-related superset
    abdomen_family = {
        "ABDOMEN",
        "UPPER_GI",
        "LOWER_GI",
        "HEPATOBILIARY",
        "UROLOGY",
        "GYNAECOLOGY",
    }
    if sel_name == "ABDOMEN" and proto_name in abdomen_family:
        return True
    if proto_name == "ABDOMEN" and sel_name in abdomen_family:
        return True

    # Thoracic family: lung–cardiac–vascular
    thorax_family = {"LUNG", "CARDIAC", "VASCULAR"}
    if sel_name in thorax_family and proto_name in thorax_family:
        return True

    # Neuro / head & neck / vascular family
    neuro_head_family = {"NEURO", "HEAD_AND_NECK", "VASCULAR"}
    if sel_name in neuro_head_family and proto_name in neuro_head_family:
        return True

    # MSK family – asymmetric behaviour:
    # - Selected MSK → show all MSK + SPINE + UPPER_LIMB + LOWER_LIMB.
    # - Selected SPINE → only SPINE.
    # - Selected UPPER_LIMB → only UPPER_LIMB.
    # - Selected LOWER_LIMB → only LOWER_LIMB.
    if sel_name == "MSK":
        if proto_name in {"MSK", "SPINE", "UPPER_LIMB", "LOWER_LIMB"}:
            return True
    elif sel_name == "SPINE":
        if proto_name == "SPINE":
            return True
    elif sel_name == "UPPER_LIMB":
        if proto_name == "UPPER_LIMB":
            return True
    elif sel_name == "LOWER_LIMB":
        if proto_name == "LOWER_LIMB":
            return True

    return False

# Helper function to fetch user and admin templates for a given modality/body part/user
def get_case_templates(
    modality,
    body_part,
    user_id=None,
    limit=8,
    indication: str | None = None,
    core_question: str | None = None,
    module: ModuleNames | None = None,
):
    """
    Return (user_templates, admin_templates) for a given modality and body_part.

    - User templates are filtered by user_id (if provided) and ordered by most recently updated.
    - Admin templates are filtered by modality/body_part and is_active=True, ordered by name.
    """
    # If absolutely no driver is provided, bail out early
    if not (modality or body_part or core_question or indication):
        return [], []


    def _modality_matches(proto_mod, selected_mod):
        if selected_mod is None:
            return True
        if proto_mod is None:
            return False
        return proto_mod == selected_mod

    # -------------------------------
    # USER TEMPLATES (structural hard gate)
    # -------------------------------

    all_user = db.session.query(UserReportTemplate)
    if user_id:
        all_user = all_user.filter_by(user_id=user_id)
    all_user = all_user.all()

    structural_user = []
    for tpl in all_user:
        mod_ok = _modality_matches(tpl.modality, modality)
        bp_ok = _body_part_matches(tpl.body_part, body_part)

        # For user templates we are deliberately more permissive:
        # - If BOTH modality and body_part are selected for the case, we prefer
        #   templates that match both, but we also allow templates that have
        #   only one dimension set and treat the missing dimension as a wildcard.
        # - If ONLY modality is selected, a user template with matching modality
        #   OR no modality specified can be used.
        # - If ONLY body_part is selected, a user template with matching body
        #   part OR no body_part specified can be used.
        # - If neither is selected, all user templates are structurally allowed.

        if modality and body_part:
            if (mod_ok and bp_ok) or (
                tpl.modality is None and bp_ok
            ) or (
                tpl.body_part is None and mod_ok
            ):
                structural_user.append(tpl)
        elif modality:
            if mod_ok or tpl.modality is None:
                structural_user.append(tpl)
        elif body_part:
            if bp_ok or tpl.body_part is None:
                structural_user.append(tpl)
        else:
            structural_user.append(tpl)


    cleaned_core = _clean_search_text(core_question) if core_question else ""
    cleaned_ind = _clean_search_text(indication) if indication else ""

    core_tokens = _tokenize(cleaned_core)
    ind_tokens = _tokenize(cleaned_ind)

    user_core_bucket = []
    user_ind_bucket = []
    user_struct_bucket = []

    for tpl in structural_user:
        blob = f"{tpl.template_name or ''} {tpl.tags or ''}".lower()
        blob_tokens = set(_tokenize(blob))

        core_hits = sum(1 for t in core_tokens if t in blob_tokens)
        ind_hits = sum(1 for t in ind_tokens if t in blob_tokens)
        # Module-based mild boost: if a module is selected and the template has the same module,
        # give it a small positive bump so these templates float higher in the list.
        module_bonus = 0
        if module is not None:
            tpl_module = getattr(tpl, "module", None)
            if tpl_module == module:
                module_bonus = 2

        score = core_hits * 4 + ind_hits + module_bonus

        if core_tokens and core_hits > 0:
            user_core_bucket.append((score, tpl))
        elif (not core_tokens) and ind_tokens and ind_hits > 0:
            user_ind_bucket.append((score, tpl))
        else:
            user_struct_bucket.append((score, tpl))

    if core_tokens and user_core_bucket:
        bucket = user_core_bucket
    elif ind_tokens and user_ind_bucket:
        bucket = user_ind_bucket
    else:
        bucket = user_struct_bucket

    bucket.sort(key=lambda x: (-x[0], x[1].template_name or ""))

    user_templates = bucket[:limit]
    user_templates = [tpl for score, tpl in user_templates]

    # -------------------------------
    # ADMIN TEMPLATES (structural hard gate)
    # -------------------------------

    all_admin = db.session.query(AdminReportTemplate).filter(
        AdminReportTemplate.is_active.is_(True)
    ).all()

    structural_admin = []
    for tpl in all_admin:
        mod_ok = _modality_matches(tpl.modality, modality)
        bp_ok = _body_part_matches(tpl.body_part, body_part)
        if modality and body_part:
            if mod_ok and bp_ok:
                structural_admin.append(tpl)
        elif modality:
            if mod_ok:
                structural_admin.append(tpl)
        elif body_part:
            if bp_ok:
                structural_admin.append(tpl)
        else:
            structural_admin.append(tpl)

    if (modality or body_part) and not structural_admin:
        return user_templates, []

    admin_core_bucket = []
    admin_ind_bucket = []
    admin_struct_bucket = []

    for tpl in structural_admin:
        blob = f"{tpl.template_name or ''} {tpl.tags or ''}".lower()
        blob_tokens = set(_tokenize(blob))

        core_hits = sum(1 for t in core_tokens if t in blob_tokens)
        ind_hits = sum(1 for t in ind_tokens if t in blob_tokens)
        # Module-based mild boost for admin templates as well:
        module_bonus = 0
        if module is not None:
            tpl_module = getattr(tpl, "module", None)
            if tpl_module == module:
                module_bonus = 2

        score = core_hits * 4 + ind_hits + module_bonus

        if core_tokens and core_hits > 0:
            admin_core_bucket.append((score, tpl))
        elif (not core_tokens) and ind_tokens and ind_hits > 0:
            admin_ind_bucket.append((score, tpl))
        else:
            admin_struct_bucket.append((score, tpl))

    if core_tokens and admin_core_bucket:
        admin_bucket = admin_core_bucket
    elif ind_tokens and admin_ind_bucket:
        admin_bucket = admin_ind_bucket
    else:
        admin_bucket = admin_struct_bucket

    admin_bucket.sort(key=lambda x: (-x[0], x[1].template_name or ""))

    admin_templates = admin_bucket[:limit]
    admin_templates = [tpl for score, tpl in admin_templates]

    return user_templates, admin_templates

# *--------------------------------------------------
# Internal helper: simple tokenizer used by case-aware search
def _tokenize(text: str | None):
    """
    Normalise and split free-text into tokens.
    Used by Case Workspace smart search for clinical indication/core question.
    """
    if not text:
        return []
    tokens = re.split(r"\W+", text.lower())
    return [t for t in tokens if len(t) >= 2]

 # *--------------------------------------------------
 # Internal helper: build AND-of-OR token filters for smarter text matching
def _build_token_filters(columns, text: str | None):
    """
    Given a list of SQLAlchemy columns and a free-text string,
    return a list of OR(...) conditions, one per token.

    Each token condition is:
        OR(col1 ILIKE '%token%', col2 ILIKE '%token%', ...)

    These conditions are intended to be AND-ed together in the query.
    """
    if not text:
        return []

    # Split on non-word characters, drop very short tokens
    tokens = [t.strip() for t in re.split(r"\W+", text) if len(t.strip()) >= 2]
    if not tokens:
        return []

    from sqlalchemy import or_ as _or

    conditions = []
    for token in tokens:
        like = f"%{token}%"
        conditions.append(
            _or(*(col.ilike(like) for col in columns))
        )
    return conditions

# *--------------------------------------------------
# Case Workspace helpers for NormalMeasurement
def get_case_measurements(
    modality: ModalityEnum | None,
    body_part: BodyPartEnum | None,
    indication: str | None = None,
    core_question: str | None = None,
    limit: int = 8,
    module: ModuleNames | None = None,
):
    """
    Smart measurement suggestions for Case Workspace (Phase 1.5, module-aware).

    Strategy:
    - Structural filters (modality/body_part) are HARD GATES.
      * If BOTH modality and body_part are provided, only measurements matching
        BOTH are allowed. If none match, return [].
      * If only modality OR only body_part is provided, structural filter uses
        that dimension alone. If none match, return [].
      * If neither is provided, all active measurements are structurally allowed.
    - Text filters (within the structurally-matched set):
      * core_question (cleaned) is the primary driver.
      * If no core match, fall back to indication (cleaned) if present.
      * If neither core nor indication match anything, fall back to the
        structural set (modality/body_part only).
    - Loose abdomen mapping: ABDOMEN also matches enums whose names contain
      'ABDOMEN' or 'GI' (e.g. UPPER_GI, LOWER_GI).
    """

    # Require at least one driver
    if not (modality or body_part or indication or core_question):
        return []


    def _modality_matches(meas_mod, selected_mod):
        if selected_mod is None:
            return True
        if meas_mod is None:
            return False
        return meas_mod == selected_mod

    # --------------------------
    # Load ALL active measurements
    # --------------------------
    all_meas = db.session.query(NormalMeasurement).filter(
        NormalMeasurement.is_active.is_(True)
    ).all()

    # --------------------------
    # Structural HARD GATE
    # --------------------------
    structural = []
    for meas in all_meas:
        mod_ok = _modality_matches(meas.modality, modality)
        bp_ok = _body_part_matches(meas.body_part, body_part)

        if modality and body_part:
            if mod_ok and bp_ok:
                structural.append(meas)
        elif modality:
            if mod_ok:
                structural.append(meas)
        elif body_part:
            if bp_ok:
                structural.append(meas)
        else:
            structural.append(meas)

    # If structural drivers exist but zero matches → empty
    if (modality or body_part) and not structural:
        return []

    if not structural:
        return []

    # --------------------------
    # TEXT CLEANING
    # --------------------------
    cleaned_core = _clean_search_text(core_question) if core_question else ""
    cleaned_ind = _clean_search_text(indication) if indication else ""

    core_tokens = _tokenize(cleaned_core)
    ind_tokens = _tokenize(cleaned_ind)

    # --------------------------
    # BUCKETS
    # --------------------------
    core_bucket: list[tuple[int, NormalMeasurement]] = []
    ind_bucket: list[tuple[int, NormalMeasurement]] = []
    struct_bucket: list[tuple[int, NormalMeasurement]] = []

    for meas in structural:
        text_blob_parts = [
            meas.name or "",
            meas.context or "",
            meas.tags or "",
            meas.reference_text or "",
        ]
        text_blob = " ".join(text_blob_parts).lower()
        blob_tokens = set(_tokenize(text_blob))

        core_hits = sum(1 for t in core_tokens if t in blob_tokens)
        ind_hits = sum(1 for t in ind_tokens if t in blob_tokens)

        # Mild module-based boost: if a module is selected and the measurement
        # belongs to the same module, push it slightly higher in the list.
        module_bonus = 0
        if module is not None:
            meas_module = getattr(meas, "module", None)
            if meas_module == module:
                module_bonus = 2

        score = core_hits * 4 + ind_hits + module_bonus

        if core_tokens and core_hits > 0:
            core_bucket.append((score, meas))
        elif (not core_tokens) and ind_tokens and ind_hits > 0:
            ind_bucket.append((score, meas))
        else:
            struct_bucket.append((score, meas))

    # --------------------------
    # Select bucket
    # --------------------------
    if core_tokens and core_bucket:
        bucket = core_bucket
    elif ind_tokens and ind_bucket:
        bucket = ind_bucket
    else:
        bucket = struct_bucket

    # --------------------------
    # SORT + RETURN
    # --------------------------
    bucket.sort(key=lambda x: (-x[0], x[1].name or ""))

    return [meas for score, meas in bucket[:limit]]

# *--------------------------------------------------
# Case Workspace helpers for Checklists (placeholder)
def get_case_checklists(
    modality: ModalityEnum | None,
    body_part: BodyPartEnum | None,
    indication: str | None = None,
    core_question: str | None = None,
    limit: int = 8,
):
    """
    Placeholder for future checklist model.

    Once a Checklist model is implemented (with modality, body_part, tags, indication, etc.),
    this function should follow the same unified Phase 1 logic as protocols/templates/measurements:
      - Structural hard gate (modality/body_part)
      - Textual buckets (core_question, indication)
      - Sorted buckets as in get_case_measurements
    For now, returns an empty list so callers can safely iterate.
    """
    return []

# Utility function to get classification/ staging system results in case workspace
def get_case_classifications(
    modality,
    body_part,
    indication: str | None = None,
    core_question: str | None = None,
    limit: int = 8,
    module: ModuleNames | None = None,
):
    """
    Return a list of ClassificationSystem objects relevant to this case (Phase 1.5, module-aware).

    Primary filter: modality and/or body_part.
    Secondary (optional) filter: free-text match on name, short_code, and tags using the core_question.
    """
    # If absolutely no driver, bail early
    if not (modality or body_part or core_question or indication):
        return []


    def _modality_matches(proto_mod, selected_mod):
        if selected_mod is None:
            return True
        if proto_mod is None:
            return False
        return proto_mod == selected_mod

    # --------------------------
    # Load ALL active classifications
    # --------------------------
    all_cls = db.session.query(ClassificationSystem).filter(
        ClassificationSystem.is_active.is_(True)
    ).all()

    # --------------------------
    # Structural HARD GATE
    # --------------------------
    structural = []
    for cs in all_cls:
        mod_ok = _modality_matches(cs.modality, modality)
        bp_ok = _body_part_matches(cs.body_part, body_part)

        if modality and body_part:
            if mod_ok and bp_ok:
                structural.append(cs)
        elif modality:
            if mod_ok:
                structural.append(cs)
        elif body_part:
            if bp_ok:
                structural.append(cs)
        else:
            structural.append(cs)

    # If structural drivers exist but zero matches → empty
    if (modality or body_part) and not structural:
        return []

    if not structural:
        return []

    # --------------------------
    # TEXT CLEANING
    # --------------------------
    cleaned_core = _clean_search_text(core_question) if core_question else ""
    cleaned_ind = _clean_search_text(indication) if indication else ""

    core_tokens = _tokenize(cleaned_core)
    ind_tokens = _tokenize(cleaned_ind)

    # --------------------------
    # BUCKETS
    # --------------------------
    core_bucket = []
    ind_bucket = []
    struct_bucket = []

    for cs in structural:
        blob = " ".join([
            cs.name or "",
            cs.short_code or "",
            cs.description or "",
            getattr(cs, "tags", "") or ""
        ]).lower()
        blob_tokens = set(_tokenize(blob))

        core_hits = sum(1 for t in core_tokens if t in blob_tokens)
        ind_hits = sum(1 for t in ind_tokens if t in blob_tokens)

        # Mild module-based boost: if a module is selected and the classification
        # belongs to the same module, push it slightly higher in the list.
        module_bonus = 0
        if module is not None:
            cs_module = getattr(cs, "module", None)
            if cs_module == module:
                module_bonus = 2

        score = core_hits * 4 + ind_hits + module_bonus

        if core_tokens and core_hits > 0:
            core_bucket.append((score, cs))
        elif (not core_tokens) and ind_tokens and ind_hits > 0:
            ind_bucket.append((score, cs))
        else:
            struct_bucket.append((score, cs))

    # --------------------------
    # Select bucket
    # --------------------------
    if core_tokens and core_bucket:
        bucket = core_bucket
    elif ind_tokens and ind_bucket:
        bucket = ind_bucket
    else:
        bucket = struct_bucket

    # --------------------------
    # SORT + RETURN
    # --------------------------
    bucket.sort(key=lambda x: (-x[0], x[1].name or ""))

    return [cs for score, cs in bucket[:limit]]


# Utility function to add default classifications:
def add_default_classification_systems():
    """
    Seed a curated set of tnm / reporting / scoring systems
    into ClassificationSystem. Idempotent: skips if name already exists.
    Assumes:
    - ClassificationSystem has at least: name, short_code, category,
        modality, body_part, description, is_active (optional).
    - category is stored as a simple string (e.g. 'staging', 'scoring').
    """

    default_systems = [
        {
            "name": "ACR TI-RADS",
            "short_code": "TI-RADS",
            "category": "scoring",
            "modality": ModalityEnum.ULTRASOUND,
            "body_part": BodyPartEnum.HEAD_AND_NECK,
            "description": "ACR Thyroid Imaging Reporting and Data System for risk stratification of thyroid nodules.",
        },
        {
            "name": "ACR BI-RADS",
            "short_code": "BI-RADS",
            "category": "scoring",
            "modality": ModalityEnum.MAMMOGRAPHY,
            "body_part": BodyPartEnum.BREAST,
            "description": "Breast Imaging Reporting and Data System for mammography, ultrasound and breast MRI.",
        },
        {
            "name": "LI-RADS",
            "short_code": "LI-RADS",
            "category": "scoring",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.HEPATOBILIARY,
            "description": "Liver Imaging Reporting and Data System for assessment of focal liver lesions in at-risk patients.",
        },
        {
            "name": "PI-RADS",
            "short_code": "PI-RADS",
            "category": "scoring",
            "modality": ModalityEnum.MRI,
            "body_part": BodyPartEnum.UROLOGY,
            "description": "Prostate Imaging Reporting and Data System for evaluation of prostate lesions on MRI.",
        },
        {
            "name": "Bosniak Renal Cyst Classification",
            "short_code": "Bosniak",
            "category": "other",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.UROLOGY,
            "description": "Bosniak classification of cystic renal lesions (I–IV) to stratify malignancy risk.",
        },
        {
            "name": "TNM – Lung Cancer",
            "short_code": "TNM Lung",
            "category": "tnm",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.LUNG,
            "description": "TNM staging for lung carcinoma (T, N, M descriptors).",
        },
        {
            "name": "TNM – Breast Cancer",
            "short_code": "TNM Breast",
            "category": "tnm",
            "modality": ModalityEnum.MAMMOGRAPHY,
            "body_part": BodyPartEnum.BREAST,
            "description": "TNM staging for breast carcinoma.",
        },
        {
            "name": "TNM – Colorectal Cancer",
            "short_code": "TNM CRC",
            "category": "tnm",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.LOWER_GI,
            "description": "TNM staging for colon and rectal carcinoma.",
        },
        {
            "name": "ASPECTS – Ischemic Stroke",
            "short_code": "ASPECTS",
            "category": "scoring",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.NEURO,
            "description": "Alberta Stroke Program Early CT Score for early ischemic change in MCA territory.",
        },
        {
            "name": "Fisher Scale – Subarachnoid Hemorrhage",
            "short_code": "Fisher",
            "category": "scoring",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.NEURO,
            "description": "Fisher grading scale for amount/distribution of subarachnoid blood on CT in SAH.",
        },
        {
            "name": "Wilkes Classification – TMJ Internal Derangement",
            "short_code": "Wilkes TMJ",
            "category": "other",
            "modality": ModalityEnum.MRI,
            "body_part": BodyPartEnum.HEAD_AND_NECK,
            "description": "Wilkes classification for temporomandibular joint internal derangement.",
        },
        {
            "name": "Lung-RADS",
            "short_code": "Lung-RADS",
            "category": "scoring",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.LUNG,
            "description": "Lung CT Screening Reporting and Data System for low-dose CT lung cancer screening.",
        },
    ]

    created = 0

    for data in default_systems:
        existing = db.session.query(ClassificationSystem).filter_by(
            name=data["name"]
        ).first()
        if existing:
            continue

        cs = ClassificationSystem(
            name=data["name"],
            short_code=data.get("short_code"),
            category=ClassificationCategoryEnum[data["category"].upper()],
            modality=data.get("modality"),
            body_part=data.get("body_part"),
            description=data.get("description"),
        )
        db.session.add(cs)
        created += 1

    if created > 0:
        db.session.commit()
        print(f"Classification systems seeded: {created}")
    else:
        print("Classification systems already present, no new entries added.")

# Utility function add default imaging protocols:
def add_default_imaging_protocols():
    """
    Seed a curated set of imaging protocols into ImagingProtocol.

    Aligned with ImagingProtocol model:
      - name (str)
      - modality (ModalityEnum)
      - body_part (BodyPartEnum)
      - indication (str)
      - is_emergency (bool)
      - uses_contrast (bool | None)
      - contrast_details (str | None)
      - technique_text (str | None)
      - parameters_json (dict | None)
      - tags (str | None)
      - is_active (bool)
    """

    default_protocols = [
        {
            "name": "CT Brain – Non-contrast (Stroke / Trauma)",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.NEURO,
            "indication": "Acute stroke, head trauma, new focal neurological deficit, severe headache with red flags.",
            "is_emergency": True,
            "uses_contrast": False,
            "contrast_details": None,
            "technique_text": (
                "Axial non-contrast CT from foramen magnum to vertex; thin slices with "
                "soft tissue and bone reconstructions; review in brain and bone windows."
            ),
            "parameters_json": {
                "slice_thickness_mm": 0.625,
                "kvp": 120,
                "pitch": "standard",
                "kernel": "soft + bone",
            },
            "tags": "ct,brain,stroke,trauma,non-contrast",
        },
        {
            "name": "CT Pulmonary Angiography (CTPA)",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.LUNG,
            "indication": "Suspected pulmonary embolism.",
            "is_emergency": True,
            "uses_contrast": True,
            "contrast_details": "70–100 mL iodinated contrast at 4–5 mL/s; bolus tracking in main pulmonary artery.",
            "technique_text": (
                "Contrast-enhanced CT chest timed to pulmonary arterial phase. "
                "Thin collimation (≤1.25 mm), coverage from lung apices to below diaphragm, "
                "breath-hold at full inspiration."
            ),
            "parameters_json": {
                "slice_thickness_mm": 1.0,
                "kvp": 100,
                "pitch": "low",
                "contrast_phase": "pulmonary_arterial",
            },
            "tags": "ct,chest,ctpa,pe,pulmonary embolism,angiography",
        },
        {
            "name": "CT Abdomen/Pelvis – Acute Abdomen",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.ABDOMEN,
            "indication": "Acute abdominal pain, suspected obstruction, perforation, sepsis or undifferentiated acute abdomen.",
            "is_emergency": True,
            "uses_contrast": True,
            "contrast_details": (
                "IV contrast in portal venous phase; consider oral contrast depending on local practice. "
                "Non-contrast phase if urolithiasis or hemorrhage suspected."
            ),
            "technique_text": (
                "CT from dome of diaphragm to symphysis pubis. Portal venous phase acquisition; "
                "reconstructions in axial, coronal, and sagittal planes."
            ),
            "parameters_json": {
                "slice_thickness_mm": 2.0,
                "kvp": 120,
                "phase": "portal_venous",
            },
            "tags": "ct,abdomen,pelvis,acute abdomen,emergency",
        },
        {
            "name": "CT KUB – Renal Colic",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.UROLOGY,
            "indication": "Suspected ureteric calculus / renal colic.",
            "is_emergency": True,
            "uses_contrast": False,
            "contrast_details": None,
            "technique_text": (
                "Low-dose non-contrast CT from upper poles of kidneys to bladder base. "
                "Thin-slice reconstructions; review in soft tissue and bone windows."
            ),
            "parameters_json": {
                "slice_thickness_mm": 2.0,
                "low_dose": True,
            },
            "tags": "ct,kub,renal,colic,stones,non-contrast",
        },
        {
            "name": "HRCT Chest – Interstitial Lung Disease",
            "modality": ModalityEnum.CT,
            "body_part": BodyPartEnum.LUNG,
            "indication": "Suspected or known interstitial lung disease.",
            "is_emergency": False,
            "uses_contrast": False,
            "contrast_details": None,
            "technique_text": (
                "High-resolution CT of chest with thin collimation, high-spatial-frequency kernel, "
                "supine and prone imaging; inspiratory ± expiratory series."
            ),
            "parameters_json": {
                "slice_thickness_mm": 1.0,
                "reconstruction_kernel": "high_resolution",
            },
            "tags": "ct,hrct,ild,interstitial lung disease",
        },
        {
            "name": "MRI Brain – Standard Protocol",
            "modality": ModalityEnum.MRI,
            "body_part": BodyPartEnum.NEURO,
            "indication": "Headache, seizures, suspected mass, demyelination, cognitive decline.",
            "is_emergency": False,
            "uses_contrast": True,
            "contrast_details": "Gadolinium-based contrast 0.1 mmol/kg if indicated (tumour, infection, inflammation).",
            "technique_text": (
                "Sagittal T1; axial T1, T2, FLAIR, DWI/ADC; GRE/SWI as indicated. "
                "Post-contrast T1-weighted imaging in axial and coronal planes where appropriate."
            ),
            "parameters_json": {
                "slice_thickness_mm": 5.0,
                "planes": ["axial", "sagittal", "coronal"],
            },
            "tags": "mri,brain,standard,headache,seizure",
        },
        {
            "name": "MRI Lumbar Spine – Radiculopathy",
            "modality": ModalityEnum.MRI,
            "body_part": BodyPartEnum.MSK,
            "indication": "Low back pain, radiculopathy, suspected canal stenosis or disc herniation.",
            "is_emergency": False,
            "uses_contrast": False,
            "contrast_details": "Contrast only if infection, tumour, or postoperative complication suspected.",
            "technique_text": (
                "Sagittal T1 and T2, sagittal STIR or T2-FS; axial T2 (and/or T1) at disc levels; "
                "coverage from T12/L1 through sacrum."
            ),
            "parameters_json": {
                "slice_thickness_mm": 3.0,
                "planes": ["sagittal", "axial"],
            },
            "tags": "mri,spine,lumbar,radiculopathy,disc",
        },
        {
            "name": "MRI Knee – Internal Derangement",
            "modality": ModalityEnum.MRI,
            "body_part": BodyPartEnum.MSK,
            "indication": "Meniscal injury, cruciate/ collateral ligament injury, cartilage lesion.",
            "is_emergency": False,
            "uses_contrast": False,
            "contrast_details": "Intra-articular contrast for MR arthrography if instability or labral-type pathology suspected.",
            "technique_text": (
                "Sagittal PD/PD-FS, coronal PD-FS, axial sequences; optional 3D isotropic sequence "
                "for multiplanar reformatting."
            ),
            "parameters_json": {
                "slice_thickness_mm": 3.0,
                "planes": ["sagittal", "coronal", "axial"],
            },
            "tags": "mri,knee,meniscus,ligament,internal derangement",
        },
        {
            "name": "X-ray Chest – PA and Lateral",
            "modality": ModalityEnum.X_RAY,
            "body_part": BodyPartEnum.LUNG,
            "indication": "General chest evaluation, pre-operative assessment, respiratory symptoms.",
            "is_emergency": False,
            "uses_contrast": False,
            "contrast_details": None,
            "technique_text": (
                "Standard PA projection at full inspiration; lateral view as indicated. "
                "Erect positioning where possible."
            ),
            "parameters_json": {
                "projections": ["PA", "lateral"],
            },
            "tags": "xray,chest,pa,lateral,pre-op",
        },
        {
            "name": "Ultrasound Abdomen – General Survey",
            "modality": ModalityEnum.ULTRASOUND,
            "body_part": BodyPartEnum.HEPATOBILIARY,
            "indication": "RUQ pain, abnormal LFTs, suspected biliary pathology, general abdominal screening.",
            "is_emergency": False,
            "uses_contrast": False,
            "contrast_details": None,
            "technique_text": (
                "Survey of liver, gallbladder, biliary tree, pancreas, spleen, kidneys, aorta and IVC; "
                "patient fasting if feasible."
            ),
            "parameters_json": {
                "transducer_frequency_MHz": "2–5",
                "approach": "subcostal and intercostal",
            },
            "tags": "usg,abdomen,hepatobiliary,screening",
        },
    ]

    created = 0

    for data in default_protocols:
        existing = db.session.query(ImagingProtocol).filter_by(
            name=data["name"]
        ).first()
        if existing:
            continue

        proto = ImagingProtocol(
            name=data["name"],
            modality=data.get("modality"),
            body_part=data.get("body_part"),
            indication=data.get("indication"),
            is_emergency=data.get("is_emergency", False),
            uses_contrast=data.get("uses_contrast"),
            contrast_details=data.get("contrast_details"),
            technique_text=data.get("technique_text"),
            parameters_json=data.get("parameters_json"),
            tags=data.get("tags"),
            is_active=True,
        )
        db.session.add(proto)
        created += 1

    if created > 0:
        db.session.commit()
        print(f"Imaging protocols seeded: {created}")
    else:
        print("Imaging protocols already present, no new entries added.")
def load_default_data():
    """Load default admin and content data from JSON."""
    json_filepath= os.path.join(basedir,'app','default_data.json')
    try:
        with open(json_filepath) as f:
            default_data = json.load(f)
            print("json file loaded")
            return default_data
    except FileNotFoundError:
        print("default_data.json not found")
        return None
    
from flask import url_for
from markupsafe import Markup


def inline_references(content, references):
    """Replace specific terms in content with inline links to references."""
    if not references:
        return content  # Return content as-is if no references are provided

    # Loop through references and replace occurrences of each term with a link
    for ref in references:
        term = ref.title  # Assuming the reference title is the term to match
        link = url_for("content_routes.view_reference", category=ref.category, display_name="References", reference_id=ref.id)
        replacement = f'<a href="{link}" class="reference-link" target="_blank">{term}</a>'
        content = content.replace(term, replacement)  # Replace term with link in content

    return Markup(content)  # Return content as HTML-safe string

def get_anonymous_user_id():
    """Returns the ID of the anonymous user, creating one if necessary."""
    user = db.session.query(User).filter_by(username="anonymous").first()
    if user:
        return user.id
    else:
        add_anonymous_user()
        user = db.session.query(User).filter_by(username="anonymous").first()
        return user.id if user else None

#function related to report generation and report managment :
# ------------------------------------------------------
# Function to clear report preview files
import os, time
def cleanup_old_previews(folder, age_minutes=15):
    now = time.time()
    for filename in os.listdir(folder):
        path = os.path.join(folder, filename)
        if os.path.isfile(path) and filename.endswith('.pdf'):
            if now - os.path.getmtime(path) > age_minutes * 60:
                try:
                    os.remove(path)
                except Exception as e:
                    print(f"Error deleting {path}: {e}")
                    
#Functions to create pdf and word report files :
from io import BytesIO
import tempfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
def create_report_template_word(data,report_type,return_path_only=False):
    # Initialize document
    doc = Document()
    
    # Hero container (unchanged)
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = title_paragraph.add_run()
    logo = logo_run.add_picture('app/static/assets/images/logo-white-bg.png', height=Inches(0.5))
    title_run = title_paragraph.add_run(" WSC - A Workstation Companion App")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(52, 58, 64)
    title_paragraph.space_after = Pt(6)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Simplify your radiology workflow with our intuitive and comprehensive tools.")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
    subtitle.space_after = Pt(2)
    
    italic_text = doc.add_paragraph()
    italic_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    italic_run = italic_text.add_run("Because every hero needs a trusty companion.")
    italic_run.italic = True
    italic_run.font.size = Pt(14)
    italic_run.font.color.rgb = RGBColor(40, 167, 69)
    italic_text.space_after = Pt(2)
    
    # Document Title (unchanged)
    title = doc.add_heading('Report Template', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 100, 0)
    template_name=data.get('template_name')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{template_name}")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(108, 117, 125)
    subtitle.space_after = Pt(1)
    # Patient Information Table
    patient_info = {
        'Name': data.get('name', 'No name provided'),
        'Gender': data.get('gender', 'Not specified'),
        'ID': data.get('patient_id', 'No ID provided'),
        'Age': data.get('age', 'Not specified'),
        'DOB': data.get('dob', 'Not specified'),
        'Location': data.get('location', 'No location provided')
    }
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for idx, (field, value) in enumerate(patient_info.items()):
        row, col = divmod(idx, 2)
        cell = table.cell(row, col)
        cell.text = f"{field}: {value}"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.runs[0].font.size = Pt(12)
    
    # Clinical information section
    doc.add_heading('Clinical Context:', level=1)
    clinical_info_text = data.get('clinical_info', 'No clinical information provided')
    clinical_info_paragraph = doc.add_paragraph(clinical_info_text if clinical_info_text else 'No clinical information provided')
    clinical_info_paragraph.runs[0].font.size = Pt(12)
    clinical_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    clinical_info_paragraph.space_after = Pt(1)
    # technical information section
    doc.add_heading('Protocol/Technique:', level=1)
    technical_info_text = data.get('technical_info', 'No technique related infomration provided')
    technical_info_paragraph = doc.add_paragraph(technical_info_text if technical_info_text else 'No technique related infomration provided')
    technical_info_paragraph.runs[0].font.size = Pt(12)
    technical_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    technical_info_paragraph.space_after=Pt(1)
    #Comparison section
    doc.add_heading('Comparisons:', level=1)
    comparison_text = data.get('comparison', 'No relevant priors available for comparison')
    comparison_paragraph = doc.add_paragraph(comparison_text if comparison_text else 'No relevant priors available for comparison')
    comparison_paragraph.runs[0].font.size = Pt(12)
    comparison_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    comparison_paragraph
    # Observations Section
    doc.add_heading('Observations:', level=1)

    # Ensure observations is always a list, with a default empty list if it's None
    observations = data.get('observations', [])

    for obs in observations:
        # Skip any None observation
        if obs is None:
            continue
    
        # Initialize defaults for each observation
        section_name = "Section"  # Default value for section
        detail = "No details provided"  # Default value for details
    
        # Process each key-value pair within the observation
        for key, value in obs.items():
            if key == 'section':
                section_name = value if value else 'Section'  # Replace empty or None with default
            elif key == 'details':
                detail = value if value else 'No details provided'  # Replace empty or None with default

        # Create a new paragraph for each observation with section and details
        obs_paragraph = doc.add_paragraph()

        # Section name in bold
        section_heading = obs_paragraph.add_run(f"{section_name}: ")
        section_heading.bold = True
        section_heading.font.size = Pt(13)
        section_heading.font.color.rgb = RGBColor(0, 102, 204)  # Light Blue
    
        # Details in regular font
        details_run = obs_paragraph.add_run(detail)
        details_run.font.size = Pt(12)
        obs_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Conclusions Section
    doc.add_heading('Conclusions:', level=1)
    conclusions_text = data.get('conclusions', 'No conclusions provided')
    conclusions_paragraph = doc.add_paragraph(conclusions_text if conclusions_text else 'No conclusions provided')
    conclusions_paragraph.runs[0].font.size = Pt(12)
    conclusions_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusions_paragraph
    # Recommendations Section
    doc.add_heading('Recommendations:', level=1)
    recommendations_text = data.get('recommendations', 'No recommendations provided')
    recommendations_paragraph = doc.add_paragraph(recommendations_text if recommendations_text else 'No recommendations provided')
    recommendations_paragraph.runs[0].font.size = Pt(12)
    recommendations_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Signature and Footer (unchanged)
    signature = doc.add_paragraph()
    signature_run = signature.add_run("\nReported and electronically signed by:\nRegistration Number:\n")
    signature_run.font.size = Pt(12)
    signature_run.bold = True
    date_run = signature.add_run(f"Dated: {datetime.now().strftime('%Y-%m-%d')}")
    date_run.font.size = Pt(12)
    date_run.italic = True

    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = 'SmartReportTemplates\nBrought to you by WSCompanion: because every Hero needs a companion'
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.runs[0].font.size = Pt(10)
    footer_paragraph.runs[0].font.color.rgb = RGBColor(107, 142, 35)
    
    # Use a temporary directory
    with tempfile.TemporaryDirectory() as temp_dir:
        word_file_path = os.path.join(temp_dir, "report_template.docx")
        doc.save(word_file_path)
        
        if return_path_only:
            # Since the temporary directory will be deleted after the 'with' block,
            # we need to copy the file if we need the path
            temp_word_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_word_file.close()
            shutil.copy(word_file_path, temp_word_file.name)
            return temp_word_file.name
        else:
            # Read the file into a BytesIO buffer
            with open(word_file_path, "rb") as word_file:
                file_stream = BytesIO(word_file.read())
            return file_stream


# Function to create pdf document form form data :from fpdf import FPDF
from reportlab.lib.pagesizes import LETTER, A4
from reportlab.lib.units import inch, mm
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def create_report_template_pdf(data, return_path_only=False):
    # Create a BytesIO buffer to hold the PDF data
    buffer = BytesIO()
    
    # Create a SimpleDocTemplate object
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Build the PDF elements
    elements = []
    
    # Add logo and title
    # Assuming the logo is in 'app/static/assets/images/logo-white-bg.png'
    logo_path = 'app/static/assets/images/logo-white-bg.png'
    try:
        logo = Image(logo_path, width=1*inch, height=1*inch)
        logo.hAlign = 'CENTER'
        elements.append(logo)
    except Exception as e:
        pass  # If logo not found, skip it
    
    title_style = ParagraphStyle(
        name='TitleStyle',
        fontSize=20,
        leading=24,
        alignment=1,  # Center alignment
        textColor=colors.HexColor('#343A40'),
        spaceAfter=6
    )
    title = Paragraph("WSC - A Workstation Companion App", title_style)
    elements.append(title)
    
    subtitle_style = ParagraphStyle(
        name='SubtitleStyle',
        fontSize=16,
        leading=18,
        alignment=1,
        textColor=colors.HexColor('#6C757D'),
        spaceAfter=2
    )
    subtitle = Paragraph(
        "Simplify your radiology workflow with our intuitive and comprehensive tools.",
        subtitle_style
    )
    elements.append(subtitle)
    
    italic_text_style = ParagraphStyle(
        name='ItalicTextStyle',
        fontSize=14,
        leading=12,
        alignment=1,
        textColor=colors.HexColor('#28A745'),
        spaceAfter=12,
        italic=True
    )
    italic_text = Paragraph("Because every hero needs a trusty companion.", italic_text_style)
    elements.append(italic_text)
    
    # Document Title
    report_title_style = ParagraphStyle(
        name='ReportTitleStyle',
        fontSize=20,
        leading=24,
        alignment=1,
        textColor=colors.HexColor('#006400'),
        spaceAfter=6
    )
    report_title = Paragraph("Report Template", report_title_style)
    elements.append(report_title)
    
    template_name = data.get('template_name', '')
    template_name_style = ParagraphStyle(
        name='TemplateNameStyle',
        fontSize=14,
        leading=18,
        alignment=1,
        textColor=colors.HexColor('#6C757D'),
        spaceAfter=12
    )
    template_name_paragraph = Paragraph(template_name, template_name_style)
    elements.append(template_name_paragraph)
    
    # Patient Information Table
    patient_info_data = [
        ['Name:', data.get('name', 'No name provided'), 'Gender:', data.get('gender', 'Not specified')],
        ['ID:', data.get('patient_id', 'No ID provided'), 'Age:', data.get('age', 'Not specified')],
        ['DOB:', data.get('dob', 'Not specified'), 'Location:', data.get('location', 'No location provided')]
    ]
    
    table_style = TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
    ])
    
    patient_table = Table(patient_info_data, colWidths=[1*inch, 2.5*inch, 1*inch, 2.5*inch])
    patient_table.setStyle(table_style)
    elements.append(patient_table)
    elements.append(Spacer(1, 12))
    
    # Sections: Clinical Context, Protocol/Technique, Comparisons, Observations, Conclusions, Recommendations
    def add_section(heading, content):
        combined_style = ParagraphStyle(
        name='CombinedStyle',
        fontSize=12,
        leading=15,
        textColor=colors.black,
        spaceAfter=12,
        spaceBefore=12
        )
        combined_text = f'<b>{heading}</b> {content if content else "No information provided"}'
        elements.append(Paragraph(combined_text, combined_style))
        
        
    # Clinical Context
    add_section('Clinical Context:', data.get('clinical_info', 'No clinical information provided'))
    
    # Protocol/Technique
    add_section('Protocol/Technique:', data.get('technical_info', 'No technique information provided'))
    
    # Comparisons
    add_section('Comparisons:', data.get('comparison', 'No relevant priors available for comparison'))
    
    # Observations
    # setting styles for observation fields :
    
    # Section name in bold and blue color
    section_style = ParagraphStyle(
        name='SectionStyle',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=4
        )
    detail_style = ParagraphStyle(
        name='DetailStyle',
        fontSize=12,
        leading=15,
        leftIndent=20,
        spaceAfter=6
        )
    heading_style = ParagraphStyle(
        name='HeadingStyle',
        fontSize=14,
        leading=20,
        textColor=colors.black,
        spaceAfter=6,
        spaceBefore=12
    )
    elements.append(Paragraph('Observations:', heading_style))
    
    # Ensure observations is always a list, with a default entry if the list is empty or None
    
    observations = data.get('observations', [])  # Default to an empty list if observations is None
    section_name = ""
    detail = ""

    for obs in observations:
        if obs is None:  # Skip if the observation itself is None
            continue

        # Initialize defaults for each observation
        section_name = "Section"
        detail = "No details provided"
    
        for key, value in obs.items():
            if key == 'section':
                section_name = value if value else 'Section'  # Handle empty or None values
            elif key == 'details':
                detail = value if value else 'No details provided'  # Handle empty or None values

        # Add section and detail to the document
        section_paragraph = Paragraph(f"{section_name}:", section_style)
        elements.append(section_paragraph)
        detail_paragraph = Paragraph(detail, detail_style)
        elements.append(detail_paragraph)

    # Conclusions
    add_section('Conclusions:', data.get('conclusions', 'No conclusions provided'))
    
    # Recommendations
    add_section('Recommendations:', data.get('recommendations', 'No recommendations provided'))
    
    # Signature and Date
    signature_style = ParagraphStyle(
        name='SignatureStyle',
        fontSize=12,
        leading=15,
        spaceAfter=12,
        spaceBefore=36
    )
    registration_style=ParagraphStyle(
        name='RegistrationStyle',
        fontSize=12,
        leading=15,
        spaceAfter=12,
        spaceBefore=8,
        italic=True
    )
    signature_text = f"Reported and electronically signed by:                  "
    registration_no="Registration Number:               "
    signature_date = f"Dated: {datetime.now().strftime('%Y-%m-%d')}"
    elements.append(Paragraph(signature_text, signature_style))
    elements.append(Paragraph(registration_no, registration_style))
    elements.append(Paragraph(signature_date, signature_style))
    
    # Footer
    def add_footer(canvas, doc):
        footer_text = 'SmartReportTemplates\nBrought to you by WSCompanion: because every Hero needs a companion'
        canvas.saveState()
        canvas.setFont('Helvetica', 10)
        canvas.setFillColor(colors.HexColor('#6B8E23'))
        width, height = doc.pagesize
        footer_lines = footer_text.split('\n')
        y = 15 * mm
        for line in footer_lines:
            canvas.drawCentredString(width / 2.0, y, line)
            y -= 12  # Move up for next line
        canvas.restoreState()
    
    # Build the PDF
    doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)
    
    buffer.seek(0)
    
    if return_path_only:
        # Save buffer to a temporary file
        temp_pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf_file.write(buffer.getvalue())
        temp_pdf_file.close()
        return temp_pdf_file.name
    else:   
        # Return the buffer for direct download
        return buffer